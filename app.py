import streamlit as st
import geopandas as gpd
import pandas as pd
import numpy as np
import tempfile
import os
import zipfile
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
from matplotlib.tri import Triangulation
import matplotlib.patches as mpatches
from matplotlib.colors import LinearSegmentedColormap
import io
from shapely.geometry import Polygon, LineString
import math
import warnings
import xml.etree.ElementTree as ET
import base64
import json
from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import geojson
import requests
import contextily as ctx
warnings.filterwarnings('ignore')


# === ESTILOS PERSONALIZADOS - VERSI√ìN PREMIUM MODERNA ===
st.markdown("""
<style>
/* === FONDO GENERAL OSCURO ELEGANTE === */
.stApp {
    background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%) !important;
    color: #ffffff !important;
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
}

/* === SIDEBAR: FONDO OSCURO ELEGANTE === */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1e293b 0%, #0f172a 100%) !important;
    border-right: 1px solid rgba(255, 255, 255, 0.1);
    box-shadow: 5px 0 25px rgba(0, 0, 0, 0.5);
}

[data-testid="stSidebar"] *,
[data-testid="stSidebar"] .stMarkdown,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stText,
[data-testid="stSidebar"] .stTitle,
[data-testid="stSidebar"] .stSubheader {
    color: #ffffff !important;
    text-shadow: none !important;
}

/* T√≠tulo del sidebar elegante */
.sidebar-title {
    font-size: 1.4em;
    font-weight: 800;
    margin: 1.5em 0 1em 0;
    text-align: center;
    padding: 14px;
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%);
    border-radius: 16px;
    color: #ffffff !important;
    box-shadow: 0 6px 20px rgba(59, 130, 246, 0.3);
    border: 1px solid rgba(255, 255, 255, 0.2);
    letter-spacing: 0.5px;
}

/* Widgets del sidebar con estilo glassmorphism */
[data-testid="stSidebar"] .stSelectbox,
[data-testid="stSidebar"] .stDateInput,
[data-testid="stSidebar"] .stSlider {
    background: rgba(255, 255, 255, 0.05);
    backdrop-filter: blur(10px);
    border-radius: 12px;
    padding: 12px;
    margin: 8px 0;
    border: 1px solid rgba(255, 255, 255, 0.1);
}

/* Inputs y selects */
[data-testid="stSidebar"] .stSelectbox div,
[data-testid="stSidebar"] .stDateInput div,
[data-testid="stSidebar"] .stSlider label {
    color: #ffffff !important;
    font-weight: 600;
    font-size: 0.95em;
}

[data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] {
    background-color: rgba(255, 255, 255, 0.1) !important;
    border: 1px solid rgba(255, 255, 255, 0.2) !important;
    color: #ffffff !important;
    border-radius: 8px;
}

[data-testid="stSidebar"] .stSlider [data-baseweb="slider"] {
    color: #ffffff !important;
}

[data-testid="stSidebar"] .stDateInput [data-baseweb="input"] {
    background-color: rgba(255, 255, 255, 0.1) !important;
    border: 1px solid rgba(255, 255, 255, 0.2) !important;
    color: #ffffff !important;
    border-radius: 8px;
}

/* Botones premium */
.stButton > button {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: white !important;
    border: none !important;
    padding: 0.8em 1.5em !important;
    border-radius: 12px !important;
    font-weight: 700 !important;
    font-size: 1em !important;
    box-shadow: 0 4px 15px rgba(59, 130, 246, 0.4) !important;
    transition: all 0.3s ease !important;
    text-transform: uppercase !important;
    letter-spacing: 0.5px !important;
}

.stButton > button:hover {
    transform: translateY(-3px) !important;
    box-shadow: 0 8px 25px rgba(59, 130, 246, 0.6) !important;
    background: linear-gradient(135deg, #4f8df8 0%, #2d5fe8 100%) !important;
}

/* === HERO BANNER PRINCIPAL CON IMAGEN === */
.hero-banner {
    background: linear-gradient(rgba(15, 23, 42, 0.9), rgba(15, 23, 42, 0.95)),
                url('https://images.unsplash.com/photo-1597981309443-6e2d2a4d9c3f?ixlib=rb-4.0.3&ixid=M3wxMjA3fDB8MHxwaG90by1wYWdlfHx8fGVufDB8fHx8fA%3D%3D&auto=format&fit=crop&w=2070&q=80') !important;
    background-size: cover !important;
    background-position: center 40% !important;
    padding: 3.5em 2em !important;
    border-radius: 24px !important;
    margin-bottom: 2.5em !important;
    box-shadow: 0 20px 40px rgba(0, 0, 0, 0.4) !important;
    border: 1px solid rgba(59, 130, 246, 0.2) !important;
    position: relative !important;
    overflow: hidden !important;
}

.hero-banner::before {
    content: '' !important;
    position: absolute !important;
    top: 0 !important;
    left: 0 !important;
    right: 0 !important;
    bottom: 0 !important;
    background: linear-gradient(45deg, rgba(59, 130, 246, 0.1), rgba(29, 78, 216, 0.05)) !important;
    z-index: 1 !important;
}

.hero-content {
    position: relative !important;
    z-index: 2 !important;
    text-align: center !important;
}

.hero-title {
    color: #ffffff !important;
    font-size: 3.2em !important;
    font-weight: 900 !important;
    margin-bottom: 0.3em !important;
    text-shadow: 0 4px 12px rgba(0, 0, 0, 0.6) !important;
    letter-spacing: -0.5px !important;
    background: linear-gradient(135deg, #ffffff 0%, #93c5fd 100%) !important;
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
    background-clip: text !important;
}

.hero-subtitle {
    color: #cbd5e1 !important;
    font-size: 1.3em !important;
    font-weight: 400 !important;
    max-width: 800px !important;
    margin: 0 auto !important;
    line-height: 1.6 !important;
}

/* === PESTA√ëAS PRINCIPALES (fuera del sidebar) - SIN CAMBIOS === */
.stTabs [data-baseweb="tab-list"] {
    background: rgba(255, 255, 255, 0.05) !important;
    backdrop-filter: blur(10px) !important;
    padding: 8px 16px !important;
    border-radius: 16px !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    margin-top: 1em !important;
    gap: 8px !important;
}

.stTabs [data-baseweb="tab"] {
    color: #94a3b8 !important;
    font-weight: 600 !important;
    padding: 12px 24px !important;
    border-radius: 12px !important;
    background: transparent !important;
    transition: all 0.3s ease !important;
    border: 1px solid transparent !important;
}

.stTabs [data-baseweb="tab"]:hover {
    color: #ffffff !important;
    background: rgba(59, 130, 246, 0.2) !important;
    border-color: rgba(59, 130, 246, 0.3) !important;
    transform: translateY(-2px) !important;
}

.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: #ffffff !important;
    font-weight: 700 !important;
    border: none !important;
    box-shadow: 0 4px 15px rgba(59, 130, 246, 0.4) !important;
}

/* === PESTA√ëAS DEL SIDEBAR: FONDO BLANCO + TEXTO NEGRO (TU REQUERIMIENTO) === */
[data-testid="stSidebar"] .stTabs [data-baseweb="tab-list"] {
    background: #ffffff !important;
    border: 1px solid #e2e8f0 !important;
    padding: 8px !important;
    border-radius: 12px !important;
    gap: 6px !important;
}

[data-testid="stSidebar"] .stTabs [data-baseweb="tab"] {
    color: #000000 !important;
    background: transparent !important;
    border-radius: 8px !important;
    padding: 8px 16px !important;
    font-weight: 600 !important;
    border: 1px solid transparent !important;
}

[data-testid="stSidebar"] .stTabs [data-baseweb="tab"]:hover {
    background: #f1f5f9 !important;
    color: #000000 !important;
    border-color: #cbd5e1 !important;
}

/* Pesta√±a activa en el sidebar: blanco con texto negro (ajustado a tu preferencia) */
[data-testid="stSidebar"] .stTabs [aria-selected="true"] {
    background: #ffffff !important;
    color: #000000 !important;
    font-weight: 700 !important;
    border: 1px solid #3b82f6 !important;
}

/* === M√âTRICAS PREMIUM === */
div[data-testid="metric-container"] {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.8), rgba(15, 23, 42, 0.9)) !important;
    backdrop-filter: blur(10px) !important;
    border-radius: 20px !important;
    padding: 24px !important;
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
    border: 1px solid rgba(59, 130, 246, 0.2) !important;
    transition: all 0.3s ease !important;
}

div[data-testid="metric-container"]:hover {
    transform: translateY(-5px) !important;
    box-shadow: 0 15px 40px rgba(59, 130, 246, 0.2) !important;
    border-color: rgba(59, 130, 246, 0.4) !important;
}

div[data-testid="metric-container"] label,
div[data-testid="metric-container"] div,
div[data-testid="metric-container"] [data-testid="stMetricValue"],
div[data-testid="metric-container"] [data-testid="stMetricLabel"] {
    color: #ffffff !important;
    font-weight: 600 !important;
}

div[data-testid="metric-container"] [data-testid="stMetricValue"] {
    font-size: 2.5em !important;
    font-weight: 800 !important;
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
    background-clip: text !important;
}

/* === GR√ÅFICOS CON ESTILO OSCURO === */
.stPlotlyChart, .stPyplot {
    background: rgba(15, 23, 42, 0.8) !important;
    backdrop-filter: blur(10px) !important;
    border-radius: 20px !important;
    padding: 20px !important;
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
    border: 1px solid rgba(59, 130, 246, 0.2) !important;
}

/* === EXPANDERS ELEGANTES === */
.streamlit-expanderHeader {
    color: #ffffff !important;
    background: rgba(30, 41, 59, 0.8) !important;
    backdrop-filter: blur(10px) !important;
    border-radius: 16px !important;
    font-weight: 700 !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    padding: 16px 20px !important;
    margin-bottom: 10px !important;
}

.streamlit-expanderContent {
    background: rgba(15, 23, 42, 0.6) !important;
    border-radius: 0 0 16px 16px !important;
    padding: 20px !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    border-top: none !important;
}

/* === TEXTOS GENERALES === */
h1, h2, h3, h4, h5, h6 {
    color: #ffffff !important;
    font-weight: 800 !important;
    margin-top: 1.5em !important;
}

p, div, span, label, li {
    color: #cbd5e1 !important;
    line-height: 1.7 !important;
}

/* === DATA FRAMES TABLAS ELEGANTES === */
.dataframe {
    background: rgba(15, 23, 42, 0.8) !important;
    backdrop-filter: blur(10px) !important;
    border-radius: 16px !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    color: #ffffff !important;
}

.dataframe th {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: #ffffff !important;
    font-weight: 700 !important;
    padding: 16px !important;
}

.dataframe td {
    color: #cbd5e1 !important;
    padding: 14px 16px !important;
    border-bottom: 1px solid rgba(255, 255, 255, 0.1) !important;
}

/* === ALERTS Y MENSAJES === */
.stAlert {
    border-radius: 16px !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    backdrop-filter: blur(10px) !important;
}

/* === SCROLLBAR PERSONALIZADA === */
::-webkit-scrollbar {
    width: 10px !important;
    height: 10px !important;
}

::-webkit-scrollbar-track {
    background: rgba(15, 23, 42, 0.8) !important;
    border-radius: 10px !important;
}

::-webkit-scrollbar-thumb {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    border-radius: 10px !important;
}

::-webkit-scrollbar-thumb:hover {
    background: linear-gradient(135deg, #4f8df8 0%, #2d5fe8 100%) !important;
}

/* === IM√ÅGENES DEL SIDEBAR === */
[data-testid="stSidebar"] img {
    border-radius: 16px !important;
    border: 2px solid rgba(59, 130, 246, 0.3) !important;
    box-shadow: 0 8px 25px rgba(0, 0, 0, 0.3) !important;
    transition: all 0.3s ease !important;
}

[data-testid="stSidebar"] img:hover {
    transform: scale(1.02) !important;
    box-shadow: 0 12px 35px rgba(59, 130, 246, 0.4) !important;
    border-color: rgba(59, 130, 246, 0.6) !important;
}

/* === TARJETAS DE CULTIVOS === */
.cultivo-card {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.9), rgba(15, 23, 42, 0.95)) !important;
    border-radius: 20px !important;
    padding: 25px !important;
    border: 1px solid rgba(59, 130, 246, 0.2) !important;
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
    transition: all 0.3s ease !important;
    height: 100% !important;
}

.cultivo-card:hover {
    transform: translateY(-8px) !important;
    box-shadow: 0 20px 40px rgba(59, 130, 246, 0.2) !important;
    border-color: rgba(59, 130, 246, 0.4) !important;
}

/* === TABLERO DE CONTROL === */
.dashboard-grid {
    display: grid !important;
    grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)) !important;
    gap: 25px !important;
    margin: 30px 0 !important;
}

.dashboard-card {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.9), rgba(15, 23, 42, 0.95)) !important;
    border-radius: 20px !important;
    padding: 25px !important;
    border: 1px solid rgba(59, 130, 246, 0.2) !important;
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.3) !important;
    transition: all 0.3s ease !important;
}

.dashboard-card:hover {
    transform: translateY(-5px) !important;
    box-shadow: 0 20px 40px rgba(59, 130, 246, 0.2) !important;
}

/* === STATS BADGES === */
.stats-badge {
    display: inline-block !important;
    padding: 6px 14px !important;
    border-radius: 50px !important;
    font-size: 0.85em !important;
    font-weight: 700 !important;
    margin: 2px !important;
}

.badge-success {
    background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
    color: white !important;
}

.badge-warning {
    background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%) !important;
    color: white !important;
}

.badge-danger {
    background: linear-gradient(135deg, #ef4444 0%, #dc2626 100%) !important;
    color: white !important;
}

.badge-info {
    background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%) !important;
    color: white !important;
}
</style>
""", unsafe_allow_html=True)


# ===== CONFIGURACI√ìN DE SAT√âLITES DISPONIBLES =====
SATELITES_DISPONIBLES = {
    'SENTINEL-2': {
        'nombre': 'Sentinel-2',
        'resolucion': '10m',
        'revisita': '5 d√≠as',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B8', 'B11'],
        'indices': ['NDVI', 'NDRE', 'GNDVI', 'OSAVI', 'MCARI'],
        'icono': 'üõ∞Ô∏è'
    },
    'LANDSAT-8': {
        'nombre': 'Landsat 8',
        'resolucion': '30m',
        'revisita': '16 d√≠as',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B6', 'B7'],
        'indices': ['NDVI', 'NDWI', 'EVI', 'SAVI', 'MSAVI'],
        'icono': 'üõ∞Ô∏è'
    },
    'DATOS_SIMULADOS': {
        'nombre': 'Datos Simulados',
        'resolucion': '10m',
        'revisita': '5 d√≠as',
        'bandas': ['B2', 'B3', 'B4', 'B5', 'B8'],
        'indices': ['NDVI', 'NDRE', 'GNDVI'],
        'icono': 'üî¨'
    }
}

# ===== CONFIGURACI√ìN =====
# PAR√ÅMETROS GEE POR CULTIVO
PARAMETROS_CULTIVOS = {
    'PALMA ACEITERA': {
        'NITROGENO': {'min': 180, 'max': 250},
        'FOSFORO': {'min': 40, 'max': 60},
        'POTASIO': {'min': 250, 'max': 350},
        'MATERIA_ORGANICA_OPTIMA': 4.0,
        'HUMEDAD_OPTIMA': 0.35,
        'NDVI_OPTIMO': 0.85,
        'NDRE_OPTIMO': 0.5
    },
    'CACAO': {
        'NITROGENO': {'min': 100, 'max': 150},
        'FOSFORO': {'min': 30, 'max': 50},
        'POTASIO': {'min': 120, 'max': 180},
        'MATERIA_ORGANICA_OPTIMA': 5.0,
        'HUMEDAD_OPTIMA': 0.3,
        'NDVI_OPTIMO': 0.75,
        'NDRE_OPTIMO': 0.4
    },
    'BANANO': {
        'NITROGENO': {'min': 200, 'max': 300},
        'FOSFORO': {'min': 50, 'max': 80},
        'POTASIO': {'min': 300, 'max': 450},
        'MATERIA_ORGANICA_OPTIMA': 3.5,
        'HUMEDAD_OPTIMA': 0.4,
        'NDVI_OPTIMO': 0.9,
        'NDRE_OPTIMO': 0.45
    },
    'CAF√â': {
        'NITROGENO': {'min': 120, 'max': 180},
        'FOSFORO': {'min': 25, 'max': 45},
        'POTASIO': {'min': 150, 'max': 220},
        'MATERIA_ORGANICA_OPTIMA': 4.5,
        'HUMEDAD_OPTIMA': 0.28,
        'NDVI_OPTIMO': 0.7,
        'NDRE_OPTIMO': 0.35
    }
}

# PAR√ÅMETROS DE TEXTURA DEL SUELO POR CULTIVO
TEXTURA_SUELO_OPTIMA = {
    'PALMA ACEITERA': {
        'textura_optima': 'Franco',
        'arena_optima': 40,
        'limo_optima': 35,
        'arcilla_optima': 25,
        'densidad_aparente_optima': 1.2,
        'porosidad_optima': 0.55
    },
    'CACAO': {
        'textura_optima': 'Franco',
        'arena_optima': 45,
        'limo_optima': 35,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.1,
        'porosidad_optima': 0.6
    },
    'BANANO': {
        'textura_optima': 'Franco',
        'arena_optima': 50,
        'limo_optima': 30,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 1.25,
        'porosidad_optima': 0.5
    },
    'CAF√â': {
        'textura_optima': 'Franco Volc√°nico',
        'arena_optima': 40,
        'limo_optima': 40,
        'arcilla_optima': 20,
        'densidad_aparente_optima': 0.9,
        'porosidad_optima': 0.65
    }
}

# CLASIFICACI√ìN DE PENDIENTES
CLASIFICACION_PENDIENTES = {
    'PLANA (0-2%)': {'min': 0, 'max': 2, 'color': '#4daf4a', 'factor_erosivo': 0.1},
    'SUAVE (2-5%)': {'min': 2, 'max': 5, 'color': '#a6d96a', 'factor_erosivo': 0.3},
    'MODERADA (5-10%)': {'min': 5, 'max': 10, 'color': '#ffffbf', 'factor_erosivo': 0.6},
    'FUERTE (10-15%)': {'min': 10, 'max': 15, 'color': '#fdae61', 'factor_erosivo': 0.8},
    'MUY FUERTE (15-25%)': {'min': 15, 'max': 25, 'color': '#f46d43', 'factor_erosivo': 0.9},
    'EXTREMA (>25%)': {'min': 25, 'max': 100, 'color': '#d73027', 'factor_erosivo': 1.0}
}

# RECOMENDACIONES POR TIPO DE TEXTURA - ACTUALIZADO A NOMENCLATURA VENEZUELA/COLOMBIA
RECOMENDACIONES_TEXTURA = {
    'Franco': {
        'propiedades': [
            "Equilibrio arena-limo-arcilla",
            "Buena aireaci√≥n y drenaje",
            "CIC intermedia-alta",
            "Retenci√≥n de agua adecuada"
        ],
        'limitantes': [
            "Puede compactarse con maquinaria pesada",
            "Erosi√≥n en pendientes si no hay cobertura"
        ],
        'manejo': [
            "Mantener coberturas vivas o muertas",
            "Evitar tr√°nsito excesivo de maquinaria",
            "Fertilizaci√≥n eficiente, sin muchas p√©rdidas",
            "Ideal para densidad est√°ndar 9 x 9 m."
        ]
    },
    'Franco arcilloso': {
        'propiedades': [
            "Mayor proporci√≥n de arcilla (25‚Äì35%)",
            "Alta retenci√≥n de agua y nutrientes",
            "Drenaje natural lento",
            "Buena fertilidad natural"
        ],
        'limitantes': [
            "Riesgo de encharcamiento",
            "Compactaci√≥n f√°cil",
            "Menor oxigenaci√≥n radicular"
        ],
        'manejo': [
            "Implementar drenajes (canales y subdrenes)",
            "Subsolado previo a siembra",
            "Incorporar materia org√°nica (raquis, compost)",
            "Fertilizaci√≥n fraccionada en lluvias intensas"
        ]
    },
    'Franco arenoso-arcilloso': {
        'propiedades': [
            "Arena 40‚Äì50%, arcilla 20‚Äì30%",
            "Buen desarrollo radicular",
            "Drenaje moderado",
            "Retenci√≥n de agua moderada-baja"
        ],
        'limitantes': [
            "Riesgo de lixiviaci√≥n de nutrientes",
            "Estr√©s h√≠drico en veranos",
            "Fertilidad moderada"
        ],
        'manejo': [
            "Uso de coberturas leguminosas",
            "Aplicar mulching (raquis, hojas)",
            "Riego suplementario en sequ√≠a",
            "Fertilizaci√≥n fraccionada con √©nfasis en K y Mg"
        ]
    }
}

# ICONOS Y COLORES POR CULTIVO
ICONOS_CULTIVOS = {
    'PALMA ACEITERA': 'üå¥',
    'CACAO': 'üç´',
    'BANANO': 'üçå',
    'CAF√â': '‚òï'
}

COLORES_CULTIVOS = {
    'PALMA ACEITERA': '#228B22',
    'CACAO': '#654321',
    'BANANO': '#FFD700',
    'CAF√â': '#8B4513'
}

# PALETAS GEE MEJORADAS
PALETAS_GEE = {
    'FERTILIDAD': ['#d73027', '#f46d43', '#fdae61', '#fee08b', '#d9ef8b', '#a6d96a', '#66bd63', '#1a9850', '#006837'],
    'NITROGENO': ['#00ff00', '#80ff00', '#ffff00', '#ff8000', '#ff0000'],
    'FOSFORO': ['#0000ff', '#4040ff', '#8080ff', '#c0c0ff', '#ffffff'],
    'POTASIO': ['#4B0082', '#6A0DAD', '#8A2BE2', '#9370DB', '#D8BFD8'],
    'TEXTURA': ['#8c510a', '#d8b365', '#f6e8c3', '#c7eae5', '#5ab4ac', '#01665e'],
    'ELEVACION': ['#006837', '#1a9850', '#66bd63', '#a6d96a', '#d9ef8b', '#ffffbf', '#fee08b', '#fdae61', '#f46d43', '#d73027'],
    'PENDIENTE': ['#4daf4a', '#a6d96a', '#ffffbf', '#fdae61', '#f46d43', '#d73027']
}

# URLs de im√°genes para sidebar
IMAGENES_CULTIVOS = {
    'PALMA ACEITERA': 'https://images.unsplash.com/photo-1597981309443-6e2d2a4d9c3f?auto=format&fit=crop&w=200&h=150&q=80',
    'CACAO': 'https://images.unsplash.com/photo-1606312619070-d48b4c652a52?auto=format&fit=crop&w=200&h=150&q=80',
    'BANANO': 'https://images.unsplash.com/photo-1587479535213-1f3c862e2f1e?auto=format&fit=crop&w=200&h=150&q=80',
    'CAF√â': 'https://images.unsplash.com/photo-1495498882177-2a843e5c2a36?auto=format&fit=crop&w=200&h=150&q=80'
}

# ===== INICIALIZACI√ìN SEGURA DE VARIABLES DE CONFIGURACI√ìN =====
nutriente = None
satelite_seleccionado = "SENTINEL-2"
indice_seleccionado = "NDVI"
fecha_inicio = datetime.now() - timedelta(days=30)
fecha_fin = datetime.now()
intervalo_curvas = 5.0
resolucion_dem = 10.0

# ===== SIDEBAR MEJORADO (INTERFAZ VISUAL) =====
with st.sidebar:
    st.markdown('<div class="sidebar-title">‚öôÔ∏è CONFIGURACI√ìN</div>', unsafe_allow_html=True)
    cultivo = st.selectbox("Cultivo:", ["PALMA ACEITERA", "CACAO", "BANANO", "CAF√â"])
    st.image(IMAGENES_CULTIVOS[cultivo], use_container_width=True)
    analisis_tipo = st.selectbox("Tipo de An√°lisis:", ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK", "AN√ÅLISIS DE TEXTURA", "AN√ÅLISIS DE CURVAS DE NIVEL"])
    if analisis_tipo == "RECOMENDACIONES NPK":
        nutriente = st.selectbox("Nutriente:", ["NITR√ìGENO", "F√ìSFORO", "POTASIO"])
    
    st.subheader("üõ∞Ô∏è Fuente de Datos Satelitales")
    satelite_seleccionado = st.selectbox(
        "Sat√©lite:",
        ["SENTINEL-2", "LANDSAT-8", "DATOS_SIMULADOS"],
        help="Selecciona la fuente de datos satelitales"
    )
    if satelite_seleccionado in SATELITES_DISPONIBLES:
        info_satelite = SATELITES_DISPONIBLES[satelite_seleccionado]
        st.info(f"""
        **{info_satelite['icono']} {info_satelite['nombre']}**
        - Resoluci√≥n: {info_satelite['resolucion']}
        - Revisita: {info_satelite['revisita']}
        - √çndices: {', '.join(info_satelite['indices'][:3])}
        """)
    
    if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
        st.subheader("üìä √çndices de Vegetaci√≥n")
        if satelite_seleccionado == "SENTINEL-2":
            indice_seleccionado = st.selectbox("√çndice:", SATELITES_DISPONIBLES['SENTINEL-2']['indices'])
        elif satelite_seleccionado == "LANDSAT-8":
            indice_seleccionado = st.selectbox("√çndice:", SATELITES_DISPONIBLES['LANDSAT-8']['indices'])
        else:
            indice_seleccionado = st.selectbox("√çndice:", SATELITES_DISPONIBLES['DATOS_SIMULADOS']['indices'])

    if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
        st.subheader("üìÖ Rango Temporal")
        fecha_fin = st.date_input("Fecha fin", datetime.now())
        fecha_inicio = st.date_input("Fecha inicio", datetime.now() - timedelta(days=30))

    st.subheader("üéØ Divisi√≥n de Parcela")
    n_divisiones = st.slider("N√∫mero de zonas de manejo:", min_value=16, max_value=48, value=32)

    if analisis_tipo == "AN√ÅLISIS DE CURVAS DE NIVEL":
        st.subheader("üèîÔ∏è Configuraci√≥n Curvas de Nivel")
        intervalo_curvas = st.slider("Intervalo entre curvas (metros):", 1.0, 20.0, 5.0, 1.0)
        resolucion_dem = st.slider("Resoluci√≥n DEM (metros):", 5.0, 50.0, 10.0, 5.0)

    st.subheader("üì§ Subir Parcela")
    uploaded_file = st.file_uploader("Subir archivo de tu parcela", type=['zip', 'kml', 'kmz'],
                                     help="Formatos aceptados: Shapefile (.zip), KML (.kml), KMZ (.kmz)")

# ===== FUNCIONES AUXILIARES - CORREGIDAS PARA EPSG:4326 =====
def validar_y_corregir_crs(gdf):
    if gdf is None or len(gdf) == 0:
        return gdf
    try:
        if gdf.crs is None:
            gdf = gdf.set_crs('EPSG:4326', inplace=False)
            st.info("‚ÑπÔ∏è Se asign√≥ EPSG:4326 al archivo (no ten√≠a CRS)")
        elif str(gdf.crs).upper() != 'EPSG:4326':
            original_crs = str(gdf.crs)
            gdf = gdf.to_crs('EPSG:4326')
            st.info(f"‚ÑπÔ∏è Transformado de {original_crs} a EPSG:4326")
        return gdf
    except Exception as e:
        st.warning(f"‚ö†Ô∏è Error al corregir CRS: {str(e)}")
        return gdf

def calcular_superficie(gdf):
    try:
        if gdf is None or len(gdf) == 0:
            return 0.0
        gdf = validar_y_corregir_crs(gdf)
        bounds = gdf.total_bounds
        if bounds[0] < -180 or bounds[2] > 180 or bounds[1] < -90 or bounds[3] > 90:
            st.warning("‚ö†Ô∏è Coordenadas fuera de rango para c√°lculo preciso de √°rea")
            area_grados2 = gdf.geometry.area.sum()
            area_m2 = area_grados2 * 111000 * 111000
            return area_m2 / 10000
        gdf_projected = gdf.to_crs('EPSG:3857')
        area_m2 = gdf_projected.geometry.area.sum()
        return area_m2 / 10000
    except Exception as e:
        try:
            return gdf.geometry.area.sum() / 10000
        except:
            return 0.0

def dividir_parcela_en_zonas(gdf, n_zonas):
    if len(gdf) == 0:
        return gdf
    gdf = validar_y_corregir_crs(gdf)
    parcela_principal = gdf.iloc[0].geometry
    bounds = parcela_principal.bounds
    minx, miny, maxx, maxy = bounds
    sub_poligonos = []
    n_cols = math.ceil(math.sqrt(n_zonas))
    n_rows = math.ceil(n_zonas / n_cols)
    width = (maxx - minx) / n_cols
    height = (maxy - miny) / n_rows
    for i in range(n_rows):
        for j in range(n_cols):
            if len(sub_poligonos) >= n_zonas:
                break
            cell_minx = minx + (j * width)
            cell_maxx = minx + ((j + 1) * width)
            cell_miny = miny + (i * height)
            cell_maxy = miny + ((i + 1) * height)
            cell_poly = Polygon([(cell_minx, cell_miny), (cell_maxx, cell_miny), (cell_maxx, cell_maxy), (cell_minx, cell_maxy)])
            intersection = parcela_principal.intersection(cell_poly)
            if not intersection.is_empty and intersection.area > 0:
                sub_poligonos.append(intersection)
    if sub_poligonos:
        nuevo_gdf = gpd.GeoDataFrame({'id_zona': range(1, len(sub_poligonos) + 1), 'geometry': sub_poligonos}, crs='EPSG:4326')
        return nuevo_gdf
    else:
        return gdf

# ===== FUNCIONES PARA CARGAR ARCHIVOS =====
def cargar_shapefile_desde_zip(zip_file):
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            with zipfile.ZipFile(zip_file, 'r') as zip_ref:
                zip_ref.extractall(tmp_dir)
            shp_files = [f for f in os.listdir(tmp_dir) if f.endswith('.shp')]
            if shp_files:
                shp_path = os.path.join(tmp_dir, shp_files[0])
                gdf = gpd.read_file(shp_path)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
            else:
                st.error("‚ùå No se encontr√≥ ning√∫n archivo .shp en el ZIP")
                return None
    except Exception as e:
        st.error(f"‚ùå Error cargando shapefile desde ZIP: {str(e)}")
        return None

def parsear_kml_manual(contenido_kml):
    try:
        root = ET.fromstring(contenido_kml)
        namespaces = {'kml': 'http://www.opengis.net/kml/2.2'}
        polygons = []
        for polygon_elem in root.findall('.//kml:Polygon', namespaces):
            coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
            if coords_elem is not None and coords_elem.text:
                coord_text = coords_elem.text.strip()
                coord_list = []
                for coord_pair in coord_text.split():
                    parts = coord_pair.split(',')
                    if len(parts) >= 2:
                        lon = float(parts[0])
                        lat = float(parts[1])
                        coord_list.append((lon, lat))
                if len(coord_list) >= 3:
                    polygons.append(Polygon(coord_list))
        if not polygons:
            for multi_geom in root.findall('.//kml:MultiGeometry', namespaces):
                for polygon_elem in multi_geom.findall('.//kml:Polygon', namespaces):
                    coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
                    if coords_elem is not None and coords_elem.text:
                        coord_text = coords_elem.text.strip()
                        coord_list = []
                        for coord_pair in coord_text.split():
                            parts = coord_pair.split(',')
                            if len(parts) >= 2:
                                lon = float(parts[0])
                                lat = float(parts[1])
                                coord_list.append((lon, lat))
                        if len(coord_list) >= 3:
                            polygons.append(Polygon(coord_list))
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        else:
            for placemark in root.findall('.//kml:Placemark', namespaces):
                for elem_name in ['Polygon', 'LineString', 'Point', 'LinearRing']:
                    elem = placemark.find(f'.//kml:{elem_name}', namespaces)
                    if elem is not None:
                        coords_elem = elem.find('.//kml:coordinates', namespaces)
                        if coords_elem is not None and coords_elem.text:
                            coord_text = coords_elem.text.strip()
                            coord_list = []
                            for coord_pair in coord_text.split():
                                parts = coord_pair.split(',')
                                if len(parts) >= 2:
                                    lon = float(parts[0])
                                    lat = float(parts[1])
                                    coord_list.append((lon, lat))
                            if len(coord_list) >= 3:
                                polygons.append(Polygon(coord_list))
                            break
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        return None
    except Exception as e:
        st.error(f"‚ùå Error parseando KML manualmente: {str(e)}")
        return None

def cargar_kml(kml_file):
    try:
        if kml_file.name.endswith('.kmz'):
            with tempfile.TemporaryDirectory() as tmp_dir:
                with zipfile.ZipFile(kml_file, 'r') as zip_ref:
                    zip_ref.extractall(tmp_dir)
                kml_files = [f for f in os.listdir(tmp_dir) if f.endswith('.kml')]
                if kml_files:
                    kml_path = os.path.join(tmp_dir, kml_files[0])
                    with open(kml_path, 'r', encoding='utf-8') as f:
                        contenido = f.read()
                    gdf = parsear_kml_manual(contenido)
                    if gdf is not None:
                        return gdf
                    else:
                        try:
                            gdf = gpd.read_file(kml_path)
                            gdf = validar_y_corregir_crs(gdf)
                            return gdf
                        except:
                            st.error("‚ùå No se pudo cargar el archivo KML/KMZ")
                            return None
                else:
                    st.error("‚ùå No se encontr√≥ ning√∫n archivo .kml en el KMZ")
                    return None
        else:
            contenido = kml_file.read().decode('utf-8')
            gdf = parsear_kml_manual(contenido)
            if gdf is not None:
                return gdf
            else:
                kml_file.seek(0)
                gdf = gpd.read_file(kml_file)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
    except Exception as e:
        st.error(f"‚ùå Error cargando archivo KML/KMZ: {str(e)}")
        return None

def cargar_archivo_parcela(uploaded_file):
    try:
        if uploaded_file.name.endswith('.zip'):
            gdf = cargar_shapefile_desde_zip(uploaded_file)
        elif uploaded_file.name.endswith(('.kml', '.kmz')):
            gdf = cargar_kml(uploaded_file)
        else:
            st.error("‚ùå Formato de archivo no soportado")
            return None
        if gdf is not None:
            gdf = validar_y_corregir_crs(gdf)
            if not gdf.geometry.geom_type.str.contains('Polygon').any():
                st.warning("‚ö†Ô∏è El archivo no contiene pol√≠gonos. Intentando extraer pol√≠gonos...")
                gdf = gdf.explode()
                gdf = gdf[gdf.geometry.geom_type.isin(['Polygon', 'MultiPolygon'])]
                if len(gdf) > 0:
                    if 'id_zona' not in gdf.columns:
                        gdf['id_zona'] = range(1, len(gdf) + 1)
                    if str(gdf.crs).upper() != 'EPSG:4326':
                        st.warning(f"‚ö†Ô∏è El archivo no pudo ser convertido a EPSG:4326. CRS actual: {gdf.crs}")
                    return gdf
                else:
                    st.error("‚ùå No se encontraron pol√≠gonos en el archivo")
                    return None
        return gdf
    except Exception as e:
        st.error(f"‚ùå Error cargando archivo: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

# ===== FUNCIONES PARA DATOS SATELITALES =====
def descargar_datos_landsat8(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    try:
        st.info(f"üîç Buscando escenas Landsat 8...")
        datos_simulados = {
            'indice': indice,
            'valor_promedio': 0.65 + np.random.normal(0, 0.1),
            'fuente': 'Landsat-8',
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'id_escena': f"LC08_{np.random.randint(1000000, 9999999)}",
            'cobertura_nubes': f"{np.random.randint(0, 15)}%",
            'resolucion': '30m'
        }
        st.success(f"‚úÖ Escena Landsat 8 encontrada: {datos_simulados['id_escena']}")
        st.info(f"‚òÅÔ∏è Cobertura de nubes: {datos_simulados['cobertura_nubes']}")
        return datos_simulados
    except Exception as e:
        st.error(f"‚ùå Error procesando Landsat 8: {str(e)}")
        return None

def descargar_datos_sentinel2(gdf, fecha_inicio, fecha_fin, indice='NDVI'):
    try:
        st.info(f"üîç Buscando escenas Sentinel-2...")
        datos_simulados = {
            'indice': indice,
            'valor_promedio': 0.72 + np.random.normal(0, 0.08),
            'fuente': 'Sentinel-2',
            'fecha': datetime.now().strftime('%Y-%m-%d'),
            'id_escena': f"S2A_{np.random.randint(1000000, 9999999)}",
            'cobertura_nubes': f"{np.random.randint(0, 10)}%",
            'resolucion': '10m'
        }
        st.success(f"‚úÖ Escena Sentinel-2 encontrada: {datos_simulados['id_escena']}")
        st.info(f"‚òÅÔ∏è Cobertura de nubes: {datos_simulados['cobertura_nubes']}")
        return datos_simulados
    except Exception as e:
        st.error(f"‚ùå Error procesando Sentinel-2: {str(e)}")
        return None

def generar_datos_simulados(gdf, cultivo, indice='NDVI'):
    st.info("üî¨ Generando datos simulados...")
    datos_simulados = {
        'indice': indice,
        'valor_promedio': PARAMETROS_CULTIVOS[cultivo]['NDVI_OPTIMO'] * 0.8 + np.random.normal(0, 0.1),
        'fuente': 'Simulaci√≥n',
        'fecha': datetime.now().strftime('%Y-%m-%d'),
        'resolucion': '10m'
    }
    st.success("‚úÖ Datos simulados generados")
    return datos_simulados

# ===== FUNCI√ìN CORREGIDA PARA OBTENER DATOS DE NASA POWER =====
def obtener_datos_nasa_power(gdf, fecha_inicio, fecha_fin):
    """
    Obtiene datos meteorol√≥gicos diarios de NASA POWER para el centroide de la parcela.
    Variables: radiaci√≥n solar (ALLSKY_SFC_SW_DWN) y viento a 2m (WS2M).
    """
    try:
        centroid = gdf.geometry.unary_union.centroid
        lat = round(centroid.y, 4)
        lon = round(centroid.x, 4)
        start = fecha_inicio.strftime("%Y%m%d")
        end = fecha_fin.strftime("%Y%m%d")
        params = {
            'parameters': 'ALLSKY_SFC_SW_DWN,WS2M,T2M,PRECTOTCORR',
            'community': 'RE',
            'longitude': lon,
            'latitude': lat,
            'start': start,
            'end': end,
            'format': 'JSON'
        }
        url = "https://power.larc.nasa.gov/api/temporal/daily/point"
        response = requests.get(url, params=params, timeout=15)
        data = response.json()
        if 'properties' not in data:
            st.warning("‚ö†Ô∏è No se obtuvieron datos de NASA POWER (fuera de rango o sin conexi√≥n).")
            return None
        series = data['properties']['parameter']
        df_power = pd.DataFrame({
            'fecha': pd.to_datetime(list(series['ALLSKY_SFC_SW_DWN'].keys())),
            'radiacion_solar': list(series['ALLSKY_SFC_SW_DWN'].values()),
            'viento_2m': list(series['WS2M'].values()),
            'temperatura': list(series['T2M'].values()),
            'precipitacion': list(series['PRECTOTCORR'].values())
        })
        df_power = df_power.replace(-999, np.nan).dropna()
        if df_power.empty:
            st.warning("‚ö†Ô∏è Datos de NASA POWER no disponibles para el per√≠odo seleccionado.")
            return None
        st.success("‚úÖ Datos meteorol√≥gicos de NASA POWER cargados.")
        return df_power
    except Exception as e:
        st.error(f"‚ùå Error al obtener datos de NASA POWER: {str(e)}")
        return None

# ===== FUNCIONES DE AN√ÅLISIS GEE =====
def calcular_indices_satelitales_gee(gdf, cultivo, datos_satelitales):
    n_poligonos = len(gdf)
    resultados = []
    gdf_centroids = gdf.copy()
    gdf_centroids['centroid'] = gdf_centroids.geometry.centroid
    gdf_centroids['x'] = gdf_centroids.centroid.x
    gdf_centroids['y'] = gdf_centroids.centroid.y
    x_coords = gdf_centroids['x'].tolist()
    y_coords = gdf_centroids['y'].tolist()
    x_min, x_max = min(x_coords), max(x_coords)
    y_min, y_max = min(y_coords), max(y_coords)
    params = PARAMETROS_CULTIVOS[cultivo]
    valor_base_satelital = datos_satelitales.get('valor_promedio', 0.6) if datos_satelitales else 0.6
    for idx, row in gdf_centroids.iterrows():
        x_norm = (row['x'] - x_min) / (x_max - x_min) if x_max != x_min else 0.5
        y_norm = (row['y'] - y_min) / (y_max - y_min) if y_max != y_min else 0.5
        patron_espacial = (x_norm * 0.6 + y_norm * 0.4)
        base_mo = params['MATERIA_ORGANICA_OPTIMA'] * 0.7
        variabilidad_mo = patron_espacial * (params['MATERIA_ORGANICA_OPTIMA'] * 0.6)
        materia_organica = base_mo + variabilidad_mo + np.random.normal(0, 0.2)
        materia_organica = max(0.5, min(8.0, materia_organica))
        base_humedad = params['HUMEDAD_OPTIMA'] * 0.8
        variabilidad_humedad = patron_espacial * (params['HUMEDAD_OPTIMA'] * 0.4)
        humedad_suelo = base_humedad + variabilidad_humedad + np.random.normal(0, 0.05)
        humedad_suelo = max(0.1, min(0.8, humedad_suelo)
)
        ndvi_base = valor_base_satelital * 0.8
        ndvi_variacion = patron_espacial * (valor_base_satelital * 0.4)
        ndvi = ndvi_base + ndvi_variacion + np.random.normal(0, 0.06)
        ndvi = max(0.1, min(0.9, ndvi))
        ndre_base = params['NDRE_OPTIMO'] * 0.7
        ndre_variacion = patron_espacial * (params['NDRE_OPTIMO'] * 0.4)
        ndre = ndre_base + ndre_variacion + np.random.normal(0, 0.04)
        ndre = max(0.05, min(0.7, ndre))
        # Calcular NDWI simulado (proxy de humedad)
        ndwi = 0.2 + np.random.normal(0, 0.08)
        ndwi = max(0, min(1, ndwi))
        npk_actual = (ndvi * 0.4) + (ndre * 0.3) + ((materia_organica / 8) * 0.2) + (humedad_suelo * 0.1)
        npk_actual = max(0, min(1, npk_actual))
        resultados.append({
            'materia_organica': round(materia_organica, 2),
            'humedad_suelo': round(humedad_suelo, 3),
            'ndvi': round(ndvi, 3),
            'ndre': round(ndre, 3),
            'ndwi': round(ndwi, 3),
            'npk_actual': round(npk_actual, 3)
        })
    return resultados

def calcular_recomendaciones_npk_gee(indices, nutriente, cultivo):
    recomendaciones = []
    params = PARAMETROS_CULTIVOS[cultivo]
    for idx in indices:
        ndre = idx['ndre']
        materia_organica = idx['materia_organica']
        humedad_suelo = idx['humedad_suelo']
        ndvi = idx['ndvi']
        if nutriente == "NITR√ìGENO":
            factor_n = ((1 - ndre) * 0.6 + (1 - ndvi) * 0.4)
            n_recomendado = (factor_n * (params['NITROGENO']['max'] - params['NITROGENO']['min']) + params['NITROGENO']['min'])
            n_recomendado = max(params['NITROGENO']['min'] * 0.8, min(params['NITROGENO']['max'] * 1.2, n_recomendado))
            recomendaciones.append(round(n_recomendado, 1))
        elif nutriente == "F√ìSFORO":
            factor_p = ((1 - (materia_organica / 8)) * 0.7 + (1 - humedad_suelo) * 0.3)
            p_recomendado = (factor_p * (params['FOSFORO']['max'] - params['FOSFORO']['min']) + params['FOSFORO']['min'])
            p_recomendado = max(params['FOSFORO']['min'] * 0.8, min(params['FOSFORO']['max'] * 1.2, p_recomendado))
            recomendaciones.append(round(p_recomendado, 1))
        else:
            factor_k = ((1 - ndre) * 0.4 + (1 - humedad_suelo) * 0.4 + (1 - (materia_organica / 8)) * 0.2)
            k_recomendado = (factor_k * (params['POTASIO']['max'] - params['POTASIO']['min']) + params['POTASIO']['min'])
            k_recomendado = max(params['POTASIO']['min'] * 0.8, min(params['POTASIO']['max'] * 1.2, k_recomendado))
            recomendaciones.append(round(k_recomendado, 1))
    return recomendaciones

# ===== FUNCIONES DE TEXTURA DEL SUELO - ACTUALIZADAS CON NUEVA NOMENCLATURA =====
def clasificar_textura_suelo(arena, limo, arcilla):
    try:
        total = arena + limo + arcilla
        if total == 0:
            return "NO_DETERMINADA"
        arena_norm = (arena / total) * 100
        limo_norm = (limo / total) * 100
        arcilla_norm = (arcilla / total) * 100
        
        # Nomenclatura actualizada Venezuela/Colombia
        if arcilla_norm >= 35:
            return "Franco arcilloso"
        elif arcilla_norm >= 25 and arcilla_norm <= 35 and arena_norm >= 20 and arena_norm <= 45:
            return "Franco arcilloso"
        elif arena_norm >= 40 and arena_norm <= 50 and arcilla_norm >= 20 and arcilla_norm <= 30:
            return "Franco arenoso-arcilloso"
        elif arena_norm >= 50 and arena_norm <= 70 and arcilla_norm >= 5 and arcilla_norm <= 20:
            return "Franco arenoso-arcilloso"
        elif arcilla_norm >= 7 and arcilla_norm <= 27 and arena_norm >= 43 and arena_norm <= 52:
            return "Franco"
        elif arena_norm >= 85:
            return "Franco arenoso-arcilloso"
        else:
            return "Franco"
    except Exception as e:
        return "NO_DETERMINADA"

def analizar_textura_suelo(gdf, cultivo):
    gdf = validar_y_corregir_crs(gdf)
    params_textura = TEXTURA_SUELO_OPTIMA[cultivo]
    zonas_gdf = gdf.copy()
    zonas_gdf['area_ha'] = 0.0
    zonas_gdf['arena'] = 0.0
    zonas_gdf['limo'] = 0.0
    zonas_gdf['arcilla'] = 0.0
    zonas_gdf['textura_suelo'] = "NO_DETERMINADA"
    areas_ha_list = []
    arena_list = []
    limo_list = []
    arcilla_list = []
    textura_list = []
    for idx, row in zonas_gdf.iterrows():
        try:
            area_gdf = gpd.GeoDataFrame({'geometry': [row.geometry]}, crs=zonas_gdf.crs)
            area_ha = calcular_superficie(area_gdf)
            if hasattr(area_ha, 'iloc'):
                area_ha = float(area_ha.iloc[0])
            elif hasattr(area_ha, '__len__') and len(area_ha) > 0:
                area_ha = float(area_ha[0])
            else:
                area_ha = float(area_ha)
            centroid = row.geometry.centroid if hasattr(row.geometry, 'centroid') else row.geometry.representative_point()
            seed_value = abs(hash(f"{centroid.x:.6f}_{centroid.y:.6f}_{cultivo}_textura")) % (2**32)
            rng = np.random.RandomState(seed_value)
            lat_norm = (centroid.y + 90) / 180 if centroid.y else 0.5
            lon_norm = (centroid.x + 180) / 360 if centroid.x else 0.5
            variabilidad_local = 0.15 + 0.7 * (lat_norm * lon_norm)
            arena_optima = params_textura['arena_optima']
            limo_optima = params_textura['limo_optima']
            arcilla_optima = params_textura['arcilla_optima']
            arena_val = max(5, min(95, rng.normal(
                arena_optima * (0.8 + 0.4 * variabilidad_local),
                arena_optima * 0.15
            )))
            limo_val = max(5, min(95, rng.normal(
                limo_optima * (0.7 + 0.6 * variabilidad_local),
                limo_optima * 0.2
            )))
            arcilla_val = max(5, min(95, rng.normal(
                arcilla_optima * (0.75 + 0.5 * variabilidad_local),
                arcilla_optima * 0.15
            )))
            total = arena_val + limo_val + arcilla_val
            arena_pct = (arena_val / total) * 100
            limo_pct = (limo_val / total) * 100
            arcilla_pct = (arcilla_val / total) * 100
            textura = clasificar_textura_suelo(arena_pct, limo_pct, arcilla_pct)
            areas_ha_list.append(area_ha)
            arena_list.append(float(arena_pct))
            limo_list.append(float(limo_pct))
            arcilla_list.append(float(arcilla_pct))
            textura_list.append(textura)
        except Exception as e:
            areas_ha_list.append(0.0)
            arena_list.append(float(params_textura['arena_optima']))
            limo_list.append(float(params_textura['limo_optima']))
            arcilla_list.append(float(params_textura['arcilla_optima']))
            textura_list.append(params_textura['textura_optima'])
    zonas_gdf['area_ha'] = areas_ha_list
    zonas_gdf['arena'] = arena_list
    zonas_gdf['limo'] = limo_list
    zonas_gdf['arcilla'] = arcilla_list
    zonas_gdf['textura_suelo'] = textura_list
    return zonas_gdf

# ===== FUNCIONES DE CURVAS DE NIVEL =====
def clasificar_pendiente(pendiente_porcentaje):
    for categoria, params in CLASIFICACION_PENDIENTES.items():
        if params['min'] <= pendiente_porcentaje < params['max']:
            return categoria, params['color']
    return "EXTREMA (>25%)", CLASIFICACION_PENDIENTES['EXTREMA (>25%)']['color']

def calcular_estadisticas_pendiente_simple(pendiente_grid):
    pendiente_flat = pendiente_grid.flatten()
    pendiente_flat = pendiente_flat[~np.isnan(pendiente_flat)]
    if len(pendiente_flat) == 0:
        return {'promedio': 0, 'min': 0, 'max': 0, 'std': 0, 'distribucion': {}}
    stats = {
        'promedio': float(np.mean(pendiente_flat)),
        'min': float(np.min(pendiente_flat)),
        'max': float(np.max(pendiente_flat)),
        'std': float(np.std(pendiente_flat)),
        'distribucion': {}
    }
    for categoria, params in CLASIFICACION_PENDIENTES.items():
        mask = (pendiente_flat >= params['min']) & (pendiente_flat < params['max'])
        stats['distribucion'][categoria] = {'porcentaje': float(np.sum(mask) / len(pendiente_flat) * 100), 'color': params['color']}
    return stats

def generar_dem_sintetico(gdf, resolucion=10.0):
    """
    Genera un DEM sint√©tico determin√≠stico basado en las coordenadas de la parcela.
    Mismo input ‚Üí mismo output siempre.
    """
    gdf = validar_y_corregir_crs(gdf)
    bounds = gdf.total_bounds
    minx, miny, maxx, maxy = bounds
    
    # Crear una semilla determin√≠stica basada en las coordenadas de la parcela
    centroid = gdf.geometry.unary_union.centroid
    # Usamos las coordenadas del centroide para crear una semilla √∫nica
    seed_value = int(centroid.x * 10000 + centroid.y * 10000) % (2**32)
    
    # Inicializar el generador aleatorio con la semilla
    rng = np.random.RandomState(seed_value)
    
    num_cells = 50
    x = np.linspace(minx, maxx, num_cells)
    y = np.linspace(miny, maxy, num_cells)
    X, Y = np.meshgrid(x, y)
    
    # Valores fijos basados en la semilla
    elevacion_base = rng.uniform(100, 300)
    slope_x = rng.uniform(-0.001, 0.001)
    slope_y = rng.uniform(-0.001, 0.001)
    relief = np.zeros_like(X)
    
    n_hills = rng.randint(2, 5)
    for _ in range(n_hills):
        hill_center_x = rng.uniform(minx, maxx)
        hill_center_y = rng.uniform(miny, maxy)
        hill_radius = rng.uniform(0.001, 0.005)
        hill_height = rng.uniform(10, 50)
        dist = np.sqrt((X - hill_center_x)**2 + (Y - hill_center_y)**2)
        relief += hill_height * np.exp(-(dist**2) / (2 * hill_radius**2))
    
    noise = rng.randn(*X.shape) * 2
    Z = elevacion_base + slope_x * (X - minx) + slope_y * (Y - miny) + relief + noise
    Z = np.maximum(Z, 50)
    
    return X, Y, Z, bounds

def calcular_pendiente_simple(X, Y, Z, resolucion=10.0):
    dy = np.gradient(Z, axis=0) / resolucion
    dx = np.gradient(Z, axis=1) / resolucion
    pendiente = np.sqrt(dx**2 + dy**2) * 100
    pendiente = np.clip(pendiente, 0, 100)
    return pendiente

def crear_mapa_pendientes_simple(X, Y, pendiente_grid, gdf_original):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))
    X_flat = X.flatten()
    Y_flat = Y.flatten()
    Z_flat = pendiente_grid.flatten()
    valid_mask = ~np.isnan(Z_flat)
    if np.sum(valid_mask) > 10:
        scatter = ax1.scatter(X_flat[valid_mask], Y_flat[valid_mask], c=Z_flat[valid_mask], cmap='RdYlGn_r', s=20, alpha=0.7, vmin=0, vmax=30)
        cbar = plt.colorbar(scatter, ax=ax1, shrink=0.8)
        cbar.set_label('Pendiente (%)')
        for porcentaje in [2, 5, 10, 15, 25]:
            mask_cat = (Z_flat[valid_mask] >= porcentaje-1) & (Z_flat[valid_mask] <= porcentaje+1)
            if np.sum(mask_cat) > 0:
                x_center = np.mean(X_flat[valid_mask][mask_cat])
                y_center = np.mean(Y_flat[valid_mask][mask_cat])
                ax1.text(x_center, y_center, f'{porcentaje}%', fontsize=8, fontweight='bold', ha='center', va='center', bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.8))
    else:
        ax1.text(0.5, 0.5, 'Datos insuficientes\npara mapa de calor', transform=ax1.transAxes, ha='center', va='center', fontsize=12)
    gdf_original.plot(ax=ax1, color='none', edgecolor='black', linewidth=2)
    ax1.set_title('Mapa de Calor de Pendientes', fontsize=12, fontweight='bold')
    ax1.set_xlabel('Longitud')
    ax1.set_ylabel('Latitud')
    ax1.grid(True, alpha=0.3)

    if np.sum(valid_mask) > 0:
        pendiente_data = Z_flat[valid_mask]
        ax2.hist(pendiente_data, bins=30, edgecolor='black', color='skyblue', alpha=0.7)
        for porcentaje, color in [(2, 'green'), (5, 'lightgreen'), (10, 'yellow'), (15, 'orange'), (25, 'red')]:
            ax2.axvline(x=porcentaje, color=color, linestyle='--', linewidth=1, alpha=0.7)
            ax2.text(porcentaje+0.5, ax2.get_ylim()[1]*0.9, f'{porcentaje}%', color=color, fontsize=8)
        stats_pendiente = calcular_estadisticas_pendiente_simple(pendiente_grid)
        stats_text = f"""
Estad√≠sticas:
‚Ä¢ M√≠nima: {stats_pendiente['min']:.1f}%
‚Ä¢ M√°xima: {stats_pendiente['max']:.1f}%
‚Ä¢ Promedio: {stats_pendiente['promedio']:.1f}%
‚Ä¢ Desviaci√≥n: {stats_pendiente['std']:.1f}%
"""
        ax2.text(0.02, 0.98, stats_text, transform=ax2.transAxes, fontsize=9, verticalalignment='top', bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.8))
        ax2.set_xlabel('Pendiente (%)')
        ax2.set_ylabel('Frecuencia')
        ax2.set_title('Distribuci√≥n de Pendientes', fontsize=12, fontweight='bold')
        ax2.grid(True, alpha=0.3)
    else:
        ax2.text(0.5, 0.5, 'Sin datos de pendiente', transform=ax2.transAxes, ha='center', va='center', fontsize=12)
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf, calcular_estadisticas_pendiente_simple(pendiente_grid)

def generar_curvas_nivel_simple(X, Y, Z, intervalo=5.0, gdf_original=None):
    curvas = []
    elevaciones = []
    try:
        if gdf_original is not None:
            poligono_principal = gdf_original.iloc[0].geometry
            bounds = poligono_principal.bounds
            centro = poligono_principal.centroid
            ancho = bounds[2] - bounds[0]
            alto = bounds[3] - bounds[1]
            radio_max = min(ancho, alto) / 2
            z_min, z_max = np.nanmin(Z), np.nanmax(Z)
            n_curvas = min(10, int((z_max - z_min) / intervalo))
            for i in range(1, n_curvas + 1):
                radio = radio_max * (i / n_curvas)
                circle = centro.buffer(radio)
                interseccion = poligono_principal.intersection(circle)
                if interseccion.geom_type == 'LineString':
                    curvas.append(interseccion)
                    elevaciones.append(z_min + (i * intervalo))
                elif interseccion.geom_type == 'MultiLineString':
                    for parte in interseccion.geoms:
                        curvas.append(parte)
                        elevaciones.append(z_min + (i * intervalo))
    except Exception as e:
        if gdf_original is not None:
            bounds = gdf_original.total_bounds
            for i in range(3):
                y = bounds[1] + (i + 1) * ((bounds[3] - bounds[1]) / 4)
                linea = LineString([(bounds[0], y), (bounds[2], y)])
                curvas.append(linea)
                elevaciones.append(100 + i * 50)
    return curvas, elevaciones

# ===== FUNCIONES DE EXPORTACI√ìN Y REPORTES - CORREGIDAS =====
def exportar_a_geojson(gdf, nombre_base="parcela"):
    try:
        gdf = validar_y_corregir_crs(gdf)
        geojson_data = gdf.to_json()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"{nombre_base}_{timestamp}.geojson"
        return geojson_data, nombre_archivo
    except Exception as e:
        st.error(f"‚ùå Error exportando a GeoJSON: {str(e)}")
        return None, None

def generar_resumen_estadisticas(gdf_analizado, analisis_tipo, cultivo, df_power=None):
    estadisticas = {}
    try:
        if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
            if 'npk_actual' in gdf_analizado.columns:
                estadisticas['√çndice NPK Promedio'] = f"{gdf_analizado['npk_actual'].mean():.3f}"
            if 'ndvi' in gdf_analizado.columns:
                estadisticas['NDVI Promedio'] = f"{gdf_analizado['ndvi'].mean():.3f}"
            if 'ndwi' in gdf_analizado.columns:
                estadisticas['NDWI Promedio'] = f"{gdf_analizado['ndwi'].mean():.3f}"
            if 'materia_organica' in gdf_analizado.columns:
                estadisticas['Materia Org√°nica Promedio'] = f"{gdf_analizado['materia_organica'].mean():.1f}%"
            # Datos de NASA POWER
            if df_power is not None:
                estadisticas['Radiaci√≥n Solar Promedio'] = f"{df_power['radiacion_solar'].mean():.1f} kWh/m¬≤/d√≠a"
                estadisticas['Velocidad Viento Promedio'] = f"{df_power['viento_2m'].mean():.2f} m/s"
                estadisticas['Precipitaci√≥n Promedio'] = f"{df_power['precipitacion'].mean():.2f} mm/d√≠a"  # ‚Üê NUEVO
        elif analisis_tipo == "AN√ÅLISIS DE TEXTURA":
            if 'arena' in gdf_analizado.columns:
                estadisticas['Arena Promedio'] = f"{gdf_analizado['arena'].mean():.1f}%"
                estadisticas['Limo Promedio'] = f"{gdf_analizado['limo'].mean():.1f}%"
                estadisticas['Arcilla Promedio'] = f"{gdf_analizado['arcilla'].mean():.1f}%"
            if 'textura_suelo' in gdf_analizado.columns:
                textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "N/D"
                estadisticas['Textura Predominante'] = textura_predominante
            if 'area_ha' in gdf_analizado.columns:
                estadisticas['√Årea Promedio por Zona'] = f"{gdf_analizado['area_ha'].mean():.2f} ha"
                if gdf_analizado['area_ha'].mean() > 0:
                    estadisticas['Coeficiente de Variaci√≥n'] = f"{(gdf_analizado['area_ha'].std() / gdf_analizado['area_ha'].mean() * 100):.1f}%"
    except Exception as e:
        st.warning(f"No se pudieron calcular algunas estad√≠sticas: {str(e)}")
    return estadisticas

def generar_recomendaciones_generales(gdf_analizado, analisis_tipo, cultivo):
    recomendaciones = []
    try:
        if analisis_tipo == "FERTILIDAD ACTUAL":
            if 'npk_actual' in gdf_analizado.columns:
                npk_promedio = gdf_analizado['npk_actual'].mean()
                if npk_promedio < 0.3:
                    recomendaciones.append("Fertilidad MUY BAJA: Se recomienda aplicaci√≥n urgente de fertilizantes balanceados")
                    recomendaciones.append("Considerar enmiendas org√°nicas para mejorar la estructura del suelo")
                elif npk_promedio < 0.5:
                    recomendaciones.append("Fertilidad BAJA: Recomendada aplicaci√≥n de fertilizantes seg√∫n an√°lisis de suelo")
                elif npk_promedio < 0.7:
                    recomendaciones.append("Fertilidad ADECUADA: Mantener pr√°cticas de manejo actuales")
                else:
                    recomendaciones.append("Fertilidad √ìPTIMA: Excelente condici√≥n, continuar con manejo actual")
        elif analisis_tipo == "AN√ÅLISIS DE TEXTURA":
            if 'textura_suelo' in gdf_analizado.columns:
                textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "N/D"
                if textura_predominante == "Franco arcilloso":
                    recomendaciones.append("Suelo franco arcilloso: Mejorar drenaje y evitar laboreo en condiciones h√∫medas")
                elif textura_predominante == "Franco arenoso-arcilloso":
                    recomendaciones.append("Suelo franco arenoso-arcilloso: Aumentar materia org√°nica y considerar riego frecuente")
                elif textura_predominante == "Franco":
                    recomendaciones.append("Textura franca: Condiciones √≥ptimas, mantener pr√°cticas de conservaci√≥n")
        # === RECOMENDACIONES POR CULTIVO ===
        if cultivo == "PALMA ACEITERA":
            recomendaciones.append("Para palma aceitera: Priorizar aplicaci√≥n de potasio en zonas con deficiencia.")
            recomendaciones.append("Evitar encharcamientos prolongados: implementar drenaje en zonas planas.")
        elif cultivo == "CACAO":
            recomendaciones.append("Para cacao: Mantener cobertura de sombra y alta materia org√°nica.")
            recomendaciones.append("Evitar laboreo intenso: sistema radicular superficial.")
        elif cultivo == "BANANO":
            recomendaciones.append("Para banano: Aplicar nitr√≥geno y potasio en ciclos cortos por alto consumo.")
            recomendaciones.append("Monitorear drenaje: sensible a anegamiento.")
        elif cultivo == "CAF√â":
            recomendaciones.append("Para caf√©: Mantener cobertura vegetal para reducir erosi√≥n en pendientes.")
            recomendaciones.append("Aplicar enmiendas org√°nicas estabilizadas para mantener pH √°cido (5.5‚Äì6.5).")
        recomendaciones.append("Realizar an√°lisis de suelo de laboratorio para validar resultados satelitales")
        recomendaciones.append("Considerar agricultura de precisi√≥n para aplicaci√≥n variable de insumos")
    except Exception as e:
        recomendaciones.append("Error generando recomendaciones espec√≠ficas")
    return recomendaciones

def limpiar_texto_para_pdf(texto):
    if not isinstance(texto, str):
        texto = str(texto)
    reemplazos = {
        '\u2022': '-',          # ‚Ä¢ ‚Üí -
        '\u2705': '[OK]',       # ‚úÖ
        '\u26A0\uFE0F': '[!]',  # ‚ö†Ô∏è
        '\u274C': '[X]',        # ‚ùå
        '\u2013': '-',          # ‚Äì ‚Üí -
        '\u2014': '--',         # ‚Äî ‚Üí --
        '\u2018': "'",          # ‚Äò
        '\u2019': "'",          # ‚Äô
        '\u201C': '"',          # "
        '\u201D': '"',          # "
        '\u2192': '->',         # ‚Üí
        '\u2190': '<-',         # ‚Üê
        '\u2265': '>=',         # ‚â•
        '\u2264': '<=',         # ‚â§
        '\u00A0': ' ',          # non-breaking space ‚Üí espacio normal
    }
    for original, reemplazo in reemplazos.items():
        texto = texto.replace(original, reemplazo)
    texto = texto.encode('latin-1', errors='replace').decode('latin-1')
    return texto

def generar_reporte_pdf(gdf_analizado, cultivo, analisis_tipo, area_total,
                        nutriente=None, satelite=None, indice=None,
                        mapa_buffer=None, estadisticas=None, recomendaciones=None):
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font('Arial', '', 12)
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, limpiar_texto_para_pdf(f'REPORTE DE AN√ÅLISIS AGR√çCOLA - {cultivo}'), 0, 1, 'C')
        pdf.set_font('Arial', '', 12)
        pdf.cell(0, 10, limpiar_texto_para_pdf(f'Tipo de An√°lisis: {analisis_tipo}'), 0, 1, 'C')
        pdf.cell(0, 10, limpiar_texto_para_pdf(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}'), 0, 1, 'C')
        pdf.ln(10)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, '1. INFORMACI√ìN GENERAL', 0, 1)
        pdf.set_font('Arial', '', 12)
        info_general = f"""Cultivo: {cultivo}
√Årea Total: {area_total:.2f} ha
Zonas Analizadas: {len(gdf_analizado)}
Tipo de An√°lisis: {analisis_tipo}"""
        if satelite:
            info_general += f"\nSat√©lite: {satelite}"
        if indice:
            info_general += f"\n√çndice: {indice}"
        if nutriente:
            info_general += f"\nNutriente Analizado: {nutriente}"
        for linea in info_general.strip().split('\n'):
            pdf.cell(0, 8, limpiar_texto_para_pdf(linea), 0, 1)
        pdf.ln(5)
        if estadisticas:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '2. ESTAD√çSTICAS PRINCIPALES', 0, 1)
            pdf.set_font('Arial', '', 12)
            for key, value in estadisticas.items():
                linea = f"- {key}: {value}"
                pdf.cell(0, 8, limpiar_texto_para_pdf(linea), 0, 1)
            pdf.ln(5)
        if mapa_buffer:
            try:
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, '3. MAPA DE RESULTADOS', 0, 1)
                temp_img_path = "temp_map.png"
                with open(temp_img_path, "wb") as f:
                    f.write(mapa_buffer.getvalue())
                pdf.image(temp_img_path, x=10, w=190)
                pdf.ln(5)
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
            except Exception as e:
                pdf.cell(0, 8, limpiar_texto_para_pdf(f"Error al incluir mapa: {str(e)[:50]}..."), 0, 1)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, '4. RESUMEN DE ZONAS', 0, 1)
        pdf.set_font('Arial', '', 10)
        if gdf_analizado is not None and not gdf_analizado.empty:
            columnas_mostrar = ['id_zona', 'area_ha']
            if 'npk_actual' in gdf_analizado.columns:
                columnas_mostrar.append('npk_actual')
            if 'valor_recomendado' in gdf_analizado.columns:
                columnas_mostrar.append('valor_recomendado')
            if 'textura_suelo' in gdf_analizado.columns:
                columnas_mostrar.append('textura_suelo')
            if 'ndwi' in gdf_analizado.columns:
                columnas_mostrar.append('ndwi')
            columnas_mostrar = [col for col in columnas_mostrar if col in gdf_analizado.columns]
            if columnas_mostrar:
                datos_tabla = [columnas_mostrar]
                for _, row in gdf_analizado.head(15).iterrows():
                    fila = []
                    for col in columnas_mostrar:
                        if col in gdf_analizado.columns:
                            valor = row[col]
                            if isinstance(valor, float):
                                if col in ['npk_actual', 'ndwi']:
                                    fila.append(f"{valor:.3f}")
                                else:
                                    fila.append(f"{valor:.2f}")
                            else:
                                fila.append(str(valor))
                        else:
                            fila.append("N/A")
                    datos_tabla.append(fila)
                col_widths = [190 // len(columnas_mostrar)] * len(columnas_mostrar)
                for fila in datos_tabla:
                    for i, item in enumerate(fila):
                        if i < len(col_widths):
                            pdf.cell(col_widths[i], 8, limpiar_texto_para_pdf(str(item)), border=1)
                    pdf.ln()
                pdf.ln(5)
        if recomendaciones:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '5. RECOMENDACIONES', 0, 1)
            pdf.set_font('Arial', '', 12)
            for rec in recomendaciones:
                linea = f"- {limpiar_texto_para_pdf(rec)}"
                pdf.multi_cell(0, 8, linea)
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, '6. METADATOS T√âCNICOS', 0, 1)
        pdf.set_font('Arial', '', 10)
        metadatos = f"""Generado por: Analizador Multi-Cultivo Satellital
Versi√≥n: 2.0
Fecha de generaci√≥n: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
Sistema de coordenadas: EPSG:4326 (WGS84)
N√∫mero de zonas: {len(gdf_analizado)}"""
        for linea in metadatos.strip().split('\n'):
            pdf.cell(0, 6, limpiar_texto_para_pdf(linea), 0, 1)
        pdf_output = BytesIO()
        pdf_output.write(pdf.output(dest='S').encode('latin-1'))
        pdf_output.seek(0)
        return pdf_output
    except Exception as e:
        st.error(f"‚ùå Error generando PDF: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

def generar_reporte_docx(gdf_analizado, cultivo, analisis_tipo, area_total,
                         nutriente=None, satelite=None, indice=None,
                         mapa_buffer=None, estadisticas=None, recomendaciones=None):
    try:
        doc = Document()
        title = doc.add_heading(f'REPORTE DE AN√ÅLISIS AGR√çCOLA - {cultivo}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph(f'Tipo de An√°lisis: {analisis_tipo}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha = doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        doc.add_heading('1. INFORMACI√ìN GENERAL', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = 'Cultivo'
        info_table.cell(0, 1).text = cultivo
        info_table.cell(1, 0).text = '√Årea Total'
        info_table.cell(1, 1).text = f'{area_total:.2f} ha'
        info_table.cell(2, 0).text = 'Zonas Analizadas'
        info_table.cell(2, 1).text = str(len(gdf_analizado))
        info_table.cell(3, 0).text = 'Tipo de An√°lisis'
        info_table.cell(3, 1).text = analisis_tipo
        row_count = 4
        if satelite:
            if row_count >= len(info_table.rows):
                info_table.add_row()
            info_table.cell(row_count, 0).text = 'Sat√©lite'
            info_table.cell(row_count, 1).text = satelite
            row_count += 1
        if indice:
            if row_count >= len(info_table.rows):
                info_table.add_row()
            info_table.cell(row_count, 0).text = '√çndice'
            info_table.cell(row_count, 1).text = indice
            row_count += 1
        if nutriente:
            if row_count >= len(info_table.rows):
                info_table.add_row()
            info_table.cell(row_count, 0).text = 'Nutriente Analizado'
            info_table.cell(row_count, 1).text = nutriente
        doc.add_paragraph()
        if estadisticas:
            doc.add_heading('2. ESTAD√çSTICAS PRINCIPALES', level=1)
            for key, value in estadisticas.items():
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f'{key}: ')
                run.bold = True
                p.add_run(str(value))
            doc.add_paragraph()
        if mapa_buffer:
            try:
                doc.add_heading('3. MAPA DE RESULTADOS', level=1)
                temp_img_path = "temp_map_docx.png"
                with open(temp_img_path, "wb") as f:
                    f.write(mapa_buffer.getvalue())
                doc.add_picture(temp_img_path, width=Inches(6.0))
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
                doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f'Error al incluir mapa: {str(e)[:50]}...')
        doc.add_heading('4. RESUMEN DE ZONAS', level=1)
        if gdf_analizado is not None and not gdf_analizado.empty:
            columnas_mostrar = ['id_zona', 'area_ha']
            if 'npk_actual' in gdf_analizado.columns:
                columnas_mostrar.append('npk_actual')
            if 'valor_recomendado' in gdf_analizado.columns:
                columnas_mostrar.append('valor_recomendado')
            if 'textura_suelo' in gdf_analizado.columns:
                columnas_mostrar.append('textura_suelo')
            if 'ndwi' in gdf_analizado.columns:
                columnas_mostrar.append('ndwi')
            columnas_mostrar = [col for col in columnas_mostrar if col in gdf_analizado.columns]
            if columnas_mostrar:
                tabla = doc.add_table(rows=1, cols=len(columnas_mostrar))
                tabla.style = 'Table Grid'
                for i, col in enumerate(columnas_mostrar):
                    tabla.cell(0, i).text = col.replace('_', ' ').upper()
                for idx, row in gdf_analizado.head(10).iterrows():
                    row_cells = tabla.add_row().cells
                    for i, col in enumerate(columnas_mostrar):
                        if col in gdf_analizado.columns:
                            valor = row[col]
                            if isinstance(valor, float):
                                if col in ['npk_actual', 'ndwi']:
                                    row_cells[i].text = f"{valor:.3f}"
                                else:
                                    row_cells[i].text = f"{valor:.2f}"
                            else:
                                row_cells[i].text = str(valor)
                        else:
                            row_cells[i].text = "N/A"
                doc.add_paragraph()
        if recomendaciones:
            doc.add_heading('5. RECOMENDACIONES', level=1)
            for rec in recomendaciones:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(rec)
        doc.add_heading('6. METADATOS T√âCNICOS', level=1)
        metadatos = [
            ('Generado por', 'Analizador Multi-Cultivo Satellital'),
            ('Versi√≥n', '2.0'),
            ('Fecha de generaci√≥n', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ('Sistema de coordenadas', 'EPSG:4326 (WGS84)'),
            ('N√∫mero de zonas', str(len(gdf_analizado)))
        ]
        for key, value in metadatos:
            p = doc.add_paragraph()
            run_key = p.add_run(f'{key}: ')
            run_key.bold = True
            p.add_run(value)
        docx_output = BytesIO()
        doc.save(docx_output)
        docx_output.seek(0)
        return docx_output
    except Exception as e:
        st.error(f"‚ùå Error generando DOCX: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

# ===== FUNCIONES DE VISUALIZACI√ìN MEJORADAS CON MAPAS ESRI =====
def crear_mapa_estatico_con_esri(gdf, titulo, columna_valor, analisis_tipo, nutriente, cultivo, satelite):
    """Crea mapa est√°tico con fondo ESRI Satellite"""
    try:
        # Convertir a Web Mercator para el mapa base
        gdf_plot = gdf.to_crs(epsg=3857)
        
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        
        if analisis_tipo == "FERTILIDAD ACTUAL":
            cmap = LinearSegmentedColormap.from_list('fertilidad_gee', PALETAS_GEE['FERTILIDAD'])
            vmin, vmax = 0, 1
        else:
            if nutriente == "NITR√ìGENO":
                cmap = LinearSegmentedColormap.from_list('nitrogeno_gee', PALETAS_GEE['NITROGENO'])
                vmin, vmax = (PARAMETROS_CULTIVOS[cultivo]['NITROGENO']['min'] * 0.8,
                              PARAMETROS_CULTIVOS[cultivo]['NITROGENO']['max'] * 1.2)
            elif nutriente == "F√ìSFORO":
                cmap = LinearSegmentedColormap.from_list('fosforo_gee', PALETAS_GEE['FOSFORO'])
                vmin, vmax = (PARAMETROS_CULTIVOS[cultivo]['FOSFORO']['min'] * 0.8,
                              PARAMETROS_CULTIVOS[cultivo]['FOSFORO']['max'] * 1.2)
            else:
                cmap = LinearSegmentedColormap.from_list('potasio_gee', PALETAS_GEE['POTASIO'])
                vmin, vmax = (PARAMETROS_CULTIVOS[cultivo]['POTASIO']['min'] * 0.8,
                              PARAMETROS_CULTIVOS[cultivo]['POTASIO']['max'] * 1.2)
        
        # Plot de las zonas con colores seg√∫n valor
        for idx, row in gdf_plot.iterrows():
            valor = row[columna_valor]
            valor_norm = (valor - vmin) / (vmax - vmin) if vmax != vmin else 0.5
            valor_norm = max(0, min(1, valor_norm))
            color = cmap(valor_norm)
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='black', linewidth=1.5, alpha=0.7)
            
            # Etiqueta de zona
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{valor:.1f}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='black', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9))
        
        # Agregar mapa base ESRI Satellite
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.7)
        except:
            st.warning("‚ö†Ô∏è No se pudo cargar el mapa base ESRI. Verifica la conexi√≥n a internet.")
        
        info_satelite = SATELITES_DISPONIBLES.get(satelite, SATELITES_DISPONIBLES['DATOS_SIMULADOS'])
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} AN√ÅLISIS GEE - {cultivo}\n'
                     f'{info_satelite["icono"]} {info_satelite["nombre"]} - {analisis_tipo}\n'
                     f'{columna_valor}',
                     fontsize=16, fontweight='bold', pad=20)
        ax.set_xlabel('Longitud')
        ax.set_ylabel('Latitud')
        ax.grid(True, alpha=0.3)
        
        # Barra de colores
        sm = plt.cm.ScalarMappable(cmap=cmap, norm=plt.Normalize(vmin=vmin, vmax=vmax))
        sm.set_array([])
        cbar = plt.colorbar(sm, ax=ax, shrink=0.8)
        cbar.set_label(columna_valor, fontsize=12, fontweight='bold')
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"‚ùå Error creando mapa con ESRI: {str(e)}")
        return None

def crear_mapa_texturas_con_esri(gdf_analizado, cultivo):
    """Crea mapa de texturas con fondo ESRI Satellite"""
    try:
        # Convertir a Web Mercator
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        
        colores_textura = {
            'Franco': '#c7eae5',
            'Franco arcilloso': '#5ab4ac',
            'Franco arenoso-arcilloso': '#f6e8c3',
            'NO_DETERMINADA': '#999999'
        }
        
        # Plot de cada zona con su color seg√∫n textura
        for idx, row in gdf_plot.iterrows():
            textura = row['textura_suelo']
            color = colores_textura.get(textura, '#999999')
            gdf_plot.iloc[[idx]].plot(ax=ax, color=color, edgecolor='black', linewidth=1.5, alpha=0.8)
            
            # Etiqueta de zona
            centroid = row.geometry.centroid
            ax.annotate(f"Z{row['id_zona']}\n{textura[:10]}", (centroid.x, centroid.y),
                        xytext=(5, 5), textcoords="offset points",
                        fontsize=8, color='black', weight='bold',
                        bbox=dict(boxstyle="round,pad=0.3", facecolor='white', alpha=0.9))
        
        # Agregar mapa base ESRI Satellite
        try:
            ctx.add_basemap(ax, source=ctx.providers.Esri.WorldImagery, alpha=0.6)
        except:
            st.warning("‚ö†Ô∏è No se pudo cargar el mapa base ESRI. Verifica la conexi√≥n a internet.")
        
        ax.set_title(f'{ICONOS_CULTIVOS[cultivo]} MAPA DE TEXTURAS - {cultivo}',
                     fontsize=16, fontweight='bold', pad=20)
        ax.set_xlabel('Longitud')
        ax.set_ylabel('Latitud')
        ax.grid(True, alpha=0.3)
        
        # Leyenda
        from matplotlib.patches import Patch
        legend_elements = [Patch(facecolor=color, edgecolor='black', label=textura)
                           for textura, color in colores_textura.items()]
        ax.legend(handles=legend_elements, title='Texturas', loc='upper left', bbox_to_anchor=(1.05, 1))
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"Error creando mapa de texturas: {str(e)}")
        return None

def crear_mapa_potencial_cosecha_calor(gdf_analizado, cultivo):
    """Crea mapa de calor moderno tipo Dazzet con fondo ESRI"""
    try:
        # Convertir a Web Mercator
        gdf_plot = gdf_analizado.to_crs(epsg=3857)
        
        fig, ax = plt.subplots(1, 1, figsize=(14, 10))
        
        # Obtener centroides y valores
        centroids = gdf_plot.geometry.centroid
        x = [c.x for c in centroids]
        y = [c.y for c in centroids]
        z = gdf_plot['potencial_cosecha'].values
        
        # Crear triangulaci√≥n para mapa de calor continuo
        triang = Triangulation(x, y)
        
        # Crear mapa de calor con degradado suave (similar a Dazzet)
        levels = np.linspace(0, 1, 50)
        contour = ax.tricontourf(
            triang, z, 
            levels=levels, 
            cmap='RdYlGn',  # Colores rojo-amarillo-verde
            alpha=0.85,      # Mayor transparencia
            antialiased=True
        )
        
        # A√±adir contornos para mejor definici√≥n
        ax.tricontour(
            triang, z, 
            levels=levels[::5],  # Cada 5 niveles
            colors='white', 
            linewidths=0.5, 
            alpha=0.3
        )
        
        # Dibujar pol√≠gonos de las zonas con borde sutil - SOLUCI√ìN CORRECTA
        # Separar edgecolor y alpha
        gdf_plot.plot(
            ax=ax, 
            color='none', 
            edgecolor='white',  # Color b√°sico
            linewidth=0.8, 
            alpha=0.4  # Transparencia aplicada al borde tambi√©n
        )
        
        # Etiquetas modernas para zonas
        for idx, row in gdf_plot.iterrows():
            centroid = row.geometry.centroid
            valor = row['potencial_cosecha']
            
            # Color de texto basado en valor
            text_color = 'white' if valor < 0.5 else 'black'
            
            # CORRECCI√ìN: Usar tupla RGBA v√°lida para facecolor del bbox
            face_alpha = 0.7 if valor > 0.5 else 0.3
            bbox_facecolor = (1, 1, 1, face_alpha)  # Tupla RGBA
            
            ax.annotate(
                f"Z{row['id_zona']}\n{valor:.2f}", 
                (centroid.x, centroid.y),
                xytext=(0, 0), 
                textcoords="offset points",
                fontsize=8, 
                color=text_color, 
                weight='bold',
                ha='center',
                va='center',
                bbox=dict(
                    boxstyle="round,pad=0.3", 
                    facecolor=bbox_facecolor,
                    edgecolor='none',
                    alpha=0.9
                )
            )
        
        # Agregar mapa base ESRI Satellite con mayor transparencia
        try:
            ctx.add_basemap(
                ax, 
                source=ctx.providers.Esri.WorldImagery, 
                alpha=0.4  # M√°s transparente para ver mejor el calor
            )
        except:
            # Fondo degradado si falla ESRI
            ax.set_facecolor('#f8f9fa')
        
        # T√≠tulo estilizado
        ax.set_title(
            f"üî• MAPA DE CALOR - POTENCIAL DE COSECHA - {cultivo}", 
            fontsize=18, 
            fontweight='bold',
            pad=20,
            color='#2c3e50'
        )
        
        ax.set_xlabel("Longitud", fontsize=12, fontweight='medium', color='#34495e')
        ax.set_ylabel("Latitud", fontsize=12, fontweight='medium', color='#34495e')
        ax.grid(True, alpha=0.2, color='gray', linestyle='--')
        
        # Barra de colores moderna
        cbar = plt.colorbar(
            contour, 
            ax=ax, 
            shrink=0.8,
            pad=0.02
        )
        cbar.set_label(
            "Potencial de Cosecha (0-1)", 
            fontsize=12, 
            fontweight='bold',
            color='#2c3e50'
        )
        cbar.ax.tick_params(labelsize=10, colors='#2c3e50')
        
        # Leyenda para puntos calientes
        zonas_calientes = gdf_plot[gdf_plot['potencial_cosecha'] > 0.7]
        if not zonas_calientes.empty:
            # Puntos amarillos brillantes para zonas calientes
            for idx, row in zonas_calientes.iterrows():
                centroid = row.geometry.centroid
                ax.plot(
                    centroid.x, centroid.y, 
                    'o',  # C√≠rculo
                    markersize=12, 
                    markeredgecolor='#f1c40f',
                    markeredgewidth=2,
                    markerfacecolor='#f39c12',
                    alpha=0.8
                )
            
            # A√±adir leyenda
            from matplotlib.lines import Line2D
            hot_spot = Line2D(
                [0], [0], 
                marker='o', 
                color='w',
                markerfacecolor='#f39c12', 
                markeredgecolor='#f1c40f',
                markersize=10, 
                markeredgewidth=2,
                label='Zona Caliente (Potencial > 0.7)'
            )
            ax.legend(
                handles=[hot_spot], 
                loc='upper right',
                framealpha=0.9,
                facecolor='white'
            )
        
        # A√±adir estad√≠sticas en esquina
        stats_text = f"""
        Estad√≠sticas:
        ‚Ä¢ Promedio: {gdf_plot['potencial_cosecha'].mean():.2f}
        ‚Ä¢ M√°ximo: {gdf_plot['potencial_cosecha'].max():.2f}
        ‚Ä¢ M√≠nimo: {gdf_plot['potencial_cosecha'].min():.2f}
        ‚Ä¢ Zonas Calientes: {len(zonas_calientes)}/{len(gdf_plot)}
        """
        
        ax.text(
            0.02, 0.98, 
            stats_text,
            transform=ax.transAxes,
            fontsize=9,
            verticalalignment='top',
            bbox=dict(
                boxstyle="round,pad=0.5", 
                facecolor='white', 
                edgecolor='#3498db',
                alpha=0.95
            )
        )
        
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=200, bbox_inches='tight', facecolor='#f8f9fa')
        buf.seek(0)
        plt.close()
        return buf
    except Exception as e:
        st.error(f"Error creando mapa de calor mejorado: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None
        
# ===== FUNCI√ìN PRINCIPAL DE AN√ÅLISIS (CORREGIDA) =====
def ejecutar_analisis(gdf, nutriente, analisis_tipo, n_divisiones, cultivo,
                      satelite=None, indice=None, fecha_inicio=None,
                      fecha_fin=None, intervalo_curvas=5.0, resolucion_dem=10.0):
    resultados = {
        'exitoso': False,
        'gdf_analizado': None,
        'mapa_buffer': None,
        'tabla_datos': None,
        'estadisticas': {},
        'recomendaciones': [],
        'area_total': 0,
        'df_power': None
    }
    try:
        gdf = validar_y_corregir_crs(gdf)
        area_total = calcular_superficie(gdf)
        resultados['area_total'] = area_total
        # === AN√ÅLISIS DE TEXTURA DEL SUELO ===
        if analisis_tipo == "AN√ÅLISIS DE TEXTURA":
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            gdf_analizado = analizar_textura_suelo(gdf_dividido, cultivo)
            resultados['gdf_analizado'] = gdf_analizado
            resultados['exitoso'] = True
            return resultados
        # === AN√ÅLISIS DE CURVAS DE NIVEL ===
        elif analisis_tipo == "AN√ÅLISIS DE CURVAS DE NIVEL":
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            resultados['gdf_analizado'] = gdf_dividido
            resultados['exitoso'] = True
            return resultados
        # === AN√ÅLISIS SATELITAL (FERTILIDAD O NPK) ===
        elif analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
            datos_satelitales = None
            if satelite == "SENTINEL-2":
                datos_satelitales = descargar_datos_sentinel2(gdf, fecha_inicio, fecha_fin, indice)
            elif satelite == "LANDSAT-8":
                datos_satelitales = descargar_datos_landsat8(gdf, fecha_inicio, fecha_fin, indice)
            else:
                datos_satelitales = generar_datos_simulados(gdf, cultivo, indice)
            gdf_dividido = dividir_parcela_en_zonas(gdf, n_divisiones)
            indices_gee = calcular_indices_satelitales_gee(gdf_dividido, cultivo, datos_satelitales)
            gdf_analizado = gdf_dividido.copy()
            for idx, indice_data in enumerate(indices_gee):
                for key, value in indice_data.items():
                    gdf_analizado.loc[gdf_analizado.index[idx], key] = value
            areas_ha_list = []
            for idx, row in gdf_analizado.iterrows():
                area_gdf = gpd.GeoDataFrame({'geometry': [row.geometry]}, crs=gdf_analizado.crs)
                area_ha = calcular_superficie(area_gdf)
                if hasattr(area_ha, 'iloc'):
                    area_ha = float(area_ha.iloc[0])
                elif hasattr(area_ha, '__len__') and len(area_ha) > 0:
                    area_ha = float(area_ha[0])
                else:
                    area_ha = float(area_ha)
                areas_ha_list.append(area_ha)
            gdf_analizado['area_ha'] = areas_ha_list
            if analisis_tipo == "RECOMENDACIONES NPK":
                recomendaciones_npk = calcular_recomendaciones_npk_gee(indices_gee, nutriente, cultivo)
                gdf_analizado['valor_recomendado'] = recomendaciones_npk
            resultados['gdf_analizado'] = gdf_analizado
            resultados['exitoso'] = True
            # === DATOS DE NASA POWER ===
            if satelite:
                df_power = obtener_datos_nasa_power(gdf, fecha_inicio, fecha_fin)
                if df_power is not None:
                    resultados['df_power'] = df_power
            return resultados
        else:
            st.error(f"Tipo de an√°lisis no soportado: {analisis_tipo}")
            return resultados
    except Exception as e:
        st.error(f"‚ùå Error en an√°lisis: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return resultados

# ===== FUNCIONES DE VISUALIZACI√ìN =====
def mostrar_resultados_textura(gdf_analizado, cultivo, area_total):
    st.subheader("üìä ESTAD√çSTICAS DE TEXTURA")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "NO_DETERMINADA"
        st.metric("üèóÔ∏è Textura Predominante", textura_predominante)
    with col2:
        avg_arena = gdf_analizado['arena'].mean()
        st.metric("üèñÔ∏è Arena Promedio", f"{avg_arena:.1f}%")
    with col3:
        avg_limo = gdf_analizado['limo'].mean()
        st.metric("üå´Ô∏è Limo Promedio", f"{avg_limo:.1f}%")
    with col4:
        avg_arcilla = gdf_analizado['arcilla'].mean()
        st.metric("üß± Arcilla Promedio", f"{avg_arcilla:.1f}%")
    st.subheader("üìà COMPOSICI√ìN GRANULOM√âTRICA")
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))
    composicion = [gdf_analizado['arena'].mean(), gdf_analizado['limo'].mean(), gdf_analizado['arcilla'].mean()]
    labels = ['Arena', 'Limo', 'Arcilla']
    colors_pie = ['#d8b365', '#f6e8c3', '#01665e']
    ax1.pie(composicion, labels=labels, colors=colors_pie, autopct='%1.1f%%', startangle=90)
    ax1.set_title('Composici√≥n Promedio del Suelo')
    textura_dist = gdf_analizado['textura_suelo'].value_counts()
    ax2.bar(textura_dist.index, textura_dist.values, color=[PALETAS_GEE['TEXTURA'][i % len(PALETAS_GEE['TEXTURA'])] for i in range(len(textura_dist))])
    ax2.set_title('Distribuci√≥n de Texturas')
    ax2.set_xlabel('Textura')
    ax2.set_ylabel('N√∫mero de Zonas')
    ax2.tick_params(axis='x', rotation=45)
    plt.tight_layout()
    st.pyplot(fig)
    
    st.subheader("üó∫Ô∏è MAPA DE TEXTURAS CON ESRI SATELLITE")
    mapa_texturas = crear_mapa_texturas_con_esri(gdf_analizado, cultivo)
    if mapa_texturas:
        st.image(mapa_texturas, use_container_width=True)
        st.download_button(
            "üì• Descargar Mapa de Texturas",
            mapa_texturas,
            f"mapa_texturas_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
            "image/png"
        )
    
    st.subheader("üìã TABLA DE RESULTADOS POR ZONA")
    columnas_textura = ['id_zona', 'area_ha', 'textura_suelo', 'arena', 'limo', 'arcilla']
    columnas_textura = [col for col in columnas_textura if col in gdf_analizado.columns]
    if columnas_textura:
        tabla_textura = gdf_analizado[columnas_textura].copy()
        tabla_textura.columns = ['Zona', '√Årea (ha)', 'Textura', 'Arena (%)', 'Limo (%)', 'Arcilla (%)']
        st.dataframe(tabla_textura)
    
    st.subheader("üí° RECOMENDACIONES DE MANEJO POR TEXTURA")
    if 'textura_suelo' in gdf_analizado.columns:
        textura_predominante = gdf_analizado['textura_suelo'].mode()[0] if len(gdf_analizado) > 0 else "NO_DETERMINADA"
        if textura_predominante in RECOMENDACIONES_TEXTURA:
            st.markdown(f"#### üèóÔ∏è **{textura_predominante.upper()}**")
            info_textura = RECOMENDACIONES_TEXTURA[textura_predominante]
            col1, col2, col3 = st.columns(3)
            with col1:
                st.markdown("**‚úÖ PROPIEDADES F√çSICAS**")
                for prop in info_textura['propiedades']:
                    st.markdown(f"‚Ä¢ {prop}")
            with col2:
                st.markdown("**‚ö†Ô∏è LIMITANTES**")
                for lim in info_textura['limitantes']:
                    st.markdown(f"‚Ä¢ {lim}")
            with col3:
                st.markdown("**üõ†Ô∏è MANEJO RECOMENDADO**")
                for man in info_textura['manejo']:
                    st.markdown(f"‚Ä¢ {man}")
    
    st.subheader("üíæ DESCARGAR RESULTADOS")
    if 'columnas_textura' in locals() and columnas_textura:
        tabla_textura = gdf_analizado[columnas_textura].copy()
        tabla_textura.columns = ['Zona', '√Årea (ha)', 'Textura', 'Arena (%)', 'Limo (%)', 'Arcilla (%)']
        csv = tabla_textura.to_csv(index=False)
        st.download_button(
            "üì• Descargar CSV con An√°lisis de Textura",
            csv,
            f"textura_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            "text/csv"
        )

def mostrar_resultados_curvas_nivel(X, Y, Z, pendiente_grid, curvas, elevaciones, gdf_original, cultivo, area_total):
    st.subheader("üìä ESTAD√çSTICAS TOPOGR√ÅFICAS")
    elevaciones_flat = Z.flatten()
    elevaciones_flat = elevaciones_flat[~np.isnan(elevaciones_flat)]
    if len(elevaciones_flat) > 0:
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            elevacion_promedio = np.mean(elevaciones_flat)
            st.metric("üèîÔ∏è Elevaci√≥n Promedio", f"{elevacion_promedio:.1f} m")
        with col2:
            rango_elevacion = np.max(elevaciones_flat) - np.min(elevaciones_flat)
            st.metric("üìè Rango de Elevaci√≥n", f"{rango_elevacion:.1f} m")
        with col3:
            mapa_pendientes, stats_pendiente = crear_mapa_pendientes_simple(X, Y, pendiente_grid, gdf_original)
            st.metric("üìê Pendiente Promedio", f"{stats_pendiente['promedio']:.1f}%")
        with col4:
            num_curvas = len(curvas) if curvas else 0
            st.metric("üîÑ N√∫mero de Curvas", f"{num_curvas}")
        st.subheader("üî• MAPA DE CALOR DE PENDIENTES")
        st.image(mapa_pendientes, use_container_width=True)
        st.download_button(
            "üì• Descargar Mapa de Pendientes",
            mapa_pendientes,
            f"mapa_pendientes_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
            "image/png"
        )
        st.subheader("‚ö†Ô∏è AN√ÅLISIS DE RIESGO DE EROSION")
        if 'stats_pendiente' in locals() and 'distribucion' in stats_pendiente:
            riesgo_total = 0
            for categoria, data in stats_pendiente['distribucion'].items():
                if categoria in CLASIFICACION_PENDIENTES:
                    riesgo_total += data['porcentaje'] * CLASIFICACION_PENDIENTES[categoria]['factor_erosivo']
            riesgo_promedio = riesgo_total / 100
            col1, col2, col3 = st.columns(3)
            with col1:
                if riesgo_promedio < 0.3:
                    st.success("‚úÖ **RIESGO BAJO**")
                    st.metric("Factor Riesgo", f"{riesgo_promedio:.2f}")
                elif riesgo_promedio < 0.6:
                    st.warning("‚ö†Ô∏è **RIESGO MODERADO**")
                    st.metric("Factor Riesgo", f"{riesgo_promedio:.2f}")
                else:
                    st.error("üö® **RIESGO ALTO**")
                    st.metric("Factor Riesgo", f"{riesgo_promedio:.2f}")
            with col2:
                area_total_ha = area_total
                porcentaje_critico = sum(data['porcentaje'] for cat, data in stats_pendiente['distribucion'].items()
                                         if cat in ['FUERTE (10-15%)', 'MUY FUERTE (15-25%)', 'EXTREMA (>25%)'])
                area_critica = area_total_ha * (porcentaje_critico / 100)
                st.metric("√Årea Cr√≠tica (>10%)", f"{area_critica:.2f} ha")
            with col3:
                porcentaje_manejable = sum(data['porcentaje'] for cat, data in stats_pendiente['distribucion'].items()
                                           if cat in ['PLANA (0-2%)', 'SUAVE (2-5%)', 'MODERADA (5-10%)'])
                area_manejable = area_total_ha * (porcentaje_manejable / 100)
                st.metric("√Årea Manejable (<10%)", f"{area_manejable:.2f} ha")
        st.subheader("üìà VISUALIZACI√ìN 3D DEL TERRENO")
        try:
            fig = plt.figure(figsize=(12, 8))
            ax = fig.add_subplot(111, projection='3d')
            surf = ax.plot_surface(X, Y, Z, cmap='terrain', alpha=0.8, linewidth=0)
            ax.set_xlabel('Longitud')
            ax.set_ylabel('Latitud')
            ax.set_zlabel('Elevaci√≥n (m)')
            ax.set_title(f'Modelo 3D del Terreno - {cultivo}')
            fig.colorbar(surf, ax=ax, shrink=0.5, aspect=5, label='Elevaci√≥n (m)')
            plt.tight_layout()
            st.pyplot(fig)
        except Exception as e:
            st.warning(f"No se pudo generar visualizaci√≥n 3D: {e}")
        st.subheader("üíæ DESCARGAR RESULTADOS")
        sample_points = []
        for i in range(0, X.shape[0], 5):
            for j in range(0, X.shape[1], 5):
                if not np.isnan(Z[i, j]):
                    sample_points.append({
                        'lat': Y[i, j],
                        'lon': X[i, j],
                        'elevacion_m': Z[i, j],
                        'pendiente_%': pendiente_grid[i, j]
                    })
        if sample_points:
            df_dem = pd.DataFrame(sample_points)
            csv = df_dem.to_csv(index=False)
            st.download_button(
                label="üìä Descargar Muestras DEM (CSV)",
                data=csv,
                file_name=f"dem_muestras_{cultivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv"
            )

# ===== INTERFAZ PRINCIPAL =====
if uploaded_file:
    with st.spinner("Cargando parcela..."):
        try:
            gdf = cargar_archivo_parcela(uploaded_file)
            if gdf is not None:
                st.success(f"‚úÖ **Parcela cargada exitosamente:** {len(gdf)} pol√≠gono(s)")
                area_total = calcular_superficie(gdf)
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**üìä INFORMACI√ìN DE LA PARCELA:**")
                    st.write(f"- Pol√≠gonos: {len(gdf)}")
                    st.write(f"- √Årea total: {area_total:.1f} ha")
                    st.write(f"- CRS: {gdf.crs}")
                    st.write(f"- Formato: {uploaded_file.name.split('.')[-1].upper()}")
                    st.write("**üìç Vista Previa:**")
                    fig, ax = plt.subplots(figsize=(8, 6))
                    gdf.plot(ax=ax, color='lightgreen', edgecolor='darkgreen', alpha=0.7)
                    ax.set_title(f"Parcela: {uploaded_file.name}")
                    ax.set_xlabel("Longitud")
                    ax.set_ylabel("Latitud")
                    ax.grid(True, alpha=0.3)
                    st.pyplot(fig)
                with col2:
                    st.write("**üéØ CONFIGURACI√ìN GEE:**")
                    st.write(f"- Cultivo: {ICONOS_CULTIVOS[cultivo]} {cultivo}")
                    st.write(f"- An√°lisis: {analisis_tipo}")
                    st.write(f"- Zonas: {n_divisiones}")
                    if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
                        st.write(f"- Sat√©lite: {SATELITES_DISPONIBLES[satelite_seleccionado]['nombre']}")
                        st.write(f"- √çndice: {indice_seleccionado}")
                        st.write(f"- Per√≠odo: {fecha_inicio} a {fecha_fin}")
                    elif analisis_tipo == "AN√ÅLISIS DE CURVAS DE NIVEL":
                        st.write(f"- Intervalo curvas: {intervalo_curvas} m")
                        st.write(f"- Resoluci√≥n DEM: {resolucion_dem} m")
                if st.button("üöÄ EJECUTAR AN√ÅLISIS COMPLETO", type="primary"):
                    resultados = None
                    if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
                        resultados = ejecutar_analisis(
                            gdf, nutriente, analisis_tipo, n_divisiones,
                            cultivo, satelite_seleccionado, indice_seleccionado,
                            fecha_inicio, fecha_fin
                        )
                    elif analisis_tipo == "AN√ÅLISIS DE CURVAS DE NIVEL":
                        resultados = ejecutar_analisis(
                            gdf, None, analisis_tipo, n_divisiones,
                            cultivo, None, None, None, None,
                            intervalo_curvas, resolucion_dem
                        )
                    else:  # AN√ÅLISIS DE TEXTURA
                        resultados = ejecutar_analisis(
                            gdf, None, analisis_tipo, n_divisiones,
                            cultivo, None, None, None, None
                        )
                    # GUARDAR RESULTADOS EN SESSION STATE
                    if resultados and resultados['exitoso']:
                        st.session_state['resultados_guardados'] = {
                            'gdf_analizado': resultados['gdf_analizado'],
                            'analisis_tipo': analisis_tipo,
                            'cultivo': cultivo,
                            'area_total': resultados['area_total'],
                            'nutriente': nutriente,
                            'satelite_seleccionado': satelite_seleccionado,
                            'indice_seleccionado': indice_seleccionado,
                            'mapa_buffer': resultados.get('mapa_buffer'),
                            'X': None,
                            'Y': None,
                            'Z': None,
                            'pendiente_grid': None,
                            'gdf_original': gdf if analisis_tipo == "AN√ÅLISIS DE CURVAS DE NIVEL" else None,
                            'df_power': resultados.get('df_power')
                        }
                        if analisis_tipo == "AN√ÅLISIS DE TEXTURA":
                            mostrar_resultados_textura(resultados['gdf_analizado'], cultivo, resultados['area_total'])
                        elif analisis_tipo == "AN√ÅLISIS DE CURVAS DE NIVEL":
                            X, Y, Z, _ = generar_dem_sintetico(gdf, resolucion_dem)
                            pendiente_grid = calcular_pendiente_simple(X, Y, Z, resolucion_dem)
                            curvas, elevaciones = generar_curvas_nivel_simple(X, Y, Z, intervalo_curvas, gdf)
                            st.session_state['resultados_guardados'].update({
                                'X': X, 'Y': Y, 'Z': Z, 'pendiente_grid': pendiente_grid
                            })
                            mostrar_resultados_curvas_nivel(X, Y, Z, pendiente_grid, curvas, elevaciones, gdf, cultivo, resultados['area_total'])
                        else:
                            # Mostrar resultados GEE
                            gdf_analizado = resultados['gdf_analizado']
                            col1, col2, col3, col4 = st.columns(4)
                            with col1:
                                st.metric("Zonas Analizadas", len(gdf_analizado))
                            with col2:
                                st.metric("√Årea Total", f"{resultados['area_total']:.1f} ha")
                            with col3:
                                if analisis_tipo == "FERTILIDAD ACTUAL":
                                    valor_prom = gdf_analizado['npk_actual'].mean()
                                    st.metric("√çndice NPK Promedio", f"{valor_prom:.3f}")
                                else:
                                    valor_prom = gdf_analizado['valor_recomendado'].mean()
                                    st.metric(f"{nutriente} Promedio", f"{valor_prom:.1f} kg/ha")
                            with col4:
                                if analisis_tipo == "FERTILIDAD ACTUAL" and gdf_analizado['npk_actual'].mean() > 0:
                                    coef_var = (gdf_analizado['npk_actual'].std() / gdf_analizado['npk_actual'].mean() * 100)
                                    st.metric("Coef. Variaci√≥n", f"{coef_var:.1f}%")
                                elif analisis_tipo == "RECOMENDACIONES NPK" and gdf_analizado['valor_recomendado'].mean() > 0:
                                    coef_var = (gdf_analizado['valor_recomendado'].std() / gdf_analizado['valor_recomendado'].mean() * 100)
                                    st.metric("Coef. Variaci√≥n", f"{coef_var:.1f}%")
                            # === DATOS DE NASA POWER ===
                            if resultados.get('df_power') is not None:
                                df_power = resultados['df_power']
                                st.subheader("üå§Ô∏è DATOS METEOROL√ìGICOS (NASA POWER)")
                                col5, col6, col7 = st.columns(3)
                                with col5:
                                    st.metric("‚òÄÔ∏è Radiaci√≥n Solar", f"{df_power['radiacion_solar'].mean():.1f} kWh/m¬≤/d√≠a")
                                with col6:
                                    st.metric("üí® Viento a 2m", f"{df_power['viento_2m'].mean():.2f} m/s")
                                with col7:
                                    st.metric("üíß NDWI Promedio", f"{gdf_analizado['ndwi'].mean():.3f}")

                                # === PESTA√ëAS CON NUEVA PESTA√ëA DE POTENCIAL DE COSECHA ===
                                tab_radiacion, tab_viento, tab_precip, tab_cosecha = st.tabs([
                                    "‚òÄÔ∏è Radiaci√≥n Solar",
                                    "üí® Velocidad del Viento",
                                    "üåßÔ∏è Precipitaci√≥n",
                                    "üî• Potencial de Cosecha"
                                ])

                                def crear_grafico_personalizado(series, titulo, ylabel, color_linea, fondo_grafico='#f8f9fa', color_texto='#2c3e50'):
                                    fig, ax = plt.subplots(figsize=(10, 4))
                                    ax.set_facecolor(fondo_grafico)
                                    fig.patch.set_facecolor(fondo_grafico)
                                    ax.plot(series.index, series.values, color=color_linea, linewidth=2.2)
                                    ax.set_title(titulo, fontsize=14, fontweight='bold', color=color_texto)
                                    ax.set_ylabel(ylabel, fontsize=12, color=color_texto)
                                    ax.set_xlabel("Fecha", fontsize=11, color=color_texto)
                                    ax.tick_params(axis='x', colors=color_texto, rotation=0)
                                    ax.tick_params(axis='y', colors=color_texto)
                                    ax.grid(True, color='#cbd5e0', linestyle='--', linewidth=0.7, alpha=0.7)
                                    for spine in ax.spines.values():
                                        spine.set_color('#cbd5e0')
                                    plt.tight_layout()
                                    return fig

                                def crear_grafico_barras_personalizado(series, titulo, ylabel, color_barra, fondo_grafico='#f8f9fa', color_texto='#2c3e50'):
                                    fig, ax = plt.subplots(figsize=(10, 4))
                                    ax.set_facecolor(fondo_grafico)
                                    fig.patch.set_facecolor(fondo_grafico)
                                    ax.bar(series.index, series.values, color=color_barra, alpha=0.85)
                                    ax.set_title(titulo, fontsize=14, fontweight='bold', color=color_texto)
                                    ax.set_ylabel(ylabel, fontsize=12, color=color_texto)
                                    ax.set_xlabel("Fecha", fontsize=11, color=color_texto)
                                    ax.tick_params(axis='x', colors=color_texto, rotation=0)
                                    ax.tick_params(axis='y', colors=color_texto)
                                    ax.grid(axis='y', color='#cbd5e0', linestyle='--', linewidth=0.7, alpha=0.7)
                                    for spine in ax.spines.values():
                                        spine.set_color('#cbd5e0')
                                    plt.tight_layout()
                                    return fig

                                # === PESTA√ëA: RADIACI√ìN SOLAR ===
                                with tab_radiacion:
                                    serie_rad = df_power.set_index('fecha')['radiacion_solar']
                                    prom_rad = serie_rad.mean()
                                    max_rad = serie_rad.max()
                                    min_rad = serie_rad.min()
                                    # Interpretaci√≥n simple
                                    if prom_rad > 5.5:
                                        interpretacion = "‚òÄÔ∏è **Alta radiaci√≥n**: Condiciones √≥ptimas para fotos√≠ntesis en cultivos tropicales."
                                    elif prom_rad > 4.0:
                                        interpretacion = "üå§Ô∏è **Radiaci√≥n moderada**: Adecuada para la mayor√≠a de cultivos, con posible limitaci√≥n en d√≠as nublados."
                                    else:
                                        interpretacion = "‚òÅÔ∏è **Radiaci√≥n baja**: Puede limitar el crecimiento; vigilar desarrollo vegetativo."

                                    col_r1, col_r2, col_r3 = st.columns(3)
                                    with col_r1:
                                        st.metric("Promedio", f"{prom_rad:.1f} kWh/m¬≤/d√≠a")
                                    with col_r2:
                                        st.metric("M√°ximo", f"{max_rad:.1f}")
                                    with col_r3:
                                        st.metric("M√≠nimo", f"{min_rad:.1f}")

                                    st.pyplot(crear_grafico_personalizado(
                                        serie_rad,
                                        "Evoluci√≥n Diaria de Radiaci√≥n Solar",
                                        "Radiaci√≥n (kWh/m¬≤/d√≠a)",
                                        color_linea='#e67e22'
                                    ))
                                    st.markdown(f"**Interpretaci√≥n agron√≥mica:** {interpretacion}")

                                # === PESTA√ëA: VIENTO ===
                                with tab_viento:
                                    serie_viento = df_power.set_index('fecha')['viento_2m']
                                    prom_viento = serie_viento.mean()
                                    max_viento = serie_viento.max()
                                    min_viento = serie_viento.min()
                                    if prom_viento < 2.0:
                                        interpretacion = "üçÉ **Viento suave**: Bajo riesgo de estr√©s mec√°nico o deshidrataci√≥n."
                                    elif prom_viento < 4.0:
                                        interpretacion = "üå¨Ô∏è **Viento moderado**: Aceptable; monitorear en etapas sensibles (floraci√≥n, fruto joven)."
                                    else:
                                        interpretacion = "üí® **Viento fuerte**: Alto riesgo de da√±o mec√°nico, aumento de evapotranspiraci√≥n y posible ca√≠da de frutos."

                                    col_w1, col_w2, col_w3 = st.columns(3)
                                    with col_w1:
                                        st.metric("Promedio", f"{prom_viento:.2f} m/s")
                                    with col_w2:
                                        st.metric("M√°ximo", f"{max_viento:.2f}")
                                    with col_w3:
                                        st.metric("M√≠nimo", f"{min_viento:.2f}")

                                    st.pyplot(crear_grafico_personalizado(
                                        serie_viento,
                                        "Evoluci√≥n Diaria de Velocidad del Viento",
                                        "Viento a 2m (m/s)",
                                        color_linea='#3498db'
                                    ))
                                    st.markdown(f"**Interpretaci√≥n agron√≥mica:** {interpretacion}")

                                # === PESTA√ëA: PRECIPITACI√ìN ===
                                with tab_precip:
                                    serie_precip = df_power.set_index('fecha')['precipitacion']
                                    prom_precip = serie_precip.mean()
                                    total_precip = serie_precip.sum()
                                    dias_lluvia = (serie_precip > 0.1).sum()
                                    if prom_precip > 8:
                                        interpretacion = "üåßÔ∏è **Precipitaci√≥n alta**: Riesgo de encharcamiento y lixiviaci√≥n de nutrientes. Asegurar drenaje."
                                    elif prom_precip > 3:
                                        interpretacion = "üíß **Precipitaci√≥n adecuada**: Condiciones h√≠dricas favorables para cultivos tropicales."
                                    else:
                                        interpretacion = "üèúÔ∏è **Precipitaci√≥n baja**: Posible d√©ficit h√≠drico; considerar riego suplementario."

                                    col_p1, col_p2, col_p3 = st.columns(3)
                                    with col_p1:
                                        st.metric("Total", f"{total_precip:.1f} mm")
                                    with col_p2:
                                        st.metric("Promedio", f"{prom_precip:.1f} mm/d√≠a")
                                    with col_p3:
                                        st.metric("D√≠as con lluvia", f"{dias_lluvia}")

                                    st.pyplot(crear_grafico_barras_personalizado(
                                        serie_precip,
                                        "Precipitaci√≥n Diaria",
                                        "Precipitaci√≥n (mm/d√≠a)",
                                        color_barra='#2ecc71'
                                    ))
                                    st.markdown(f"**Interpretaci√≥n agron√≥mica:** {interpretacion}")

                                # === PESTA√ëA: POTENCIAL DE COSECHA ===
                                with tab_cosecha:
                                    st.subheader("üî• An√°lisis de Potencial de Cosecha - Puntos Calientes")
                                    st.markdown("""
                                    **Metodolog√≠a:**
                                    - Se integran datos de fertilidad (NPK), radiaci√≥n solar, humedad (NDWI) y estr√©s por viento
                                    - Las zonas con valores m√°s altos (rojo/amarillo) son los **puntos calientes** para mejor cosecha
                                    - Los c√≠rculos amarillos marcan zonas con potencial >0.7
                                    """)

                                    # --- Paso 1: Agregar datos meteorol√≥gicos promedio a cada zona ---
                                    rad_prom = df_power['radiacion_solar'].mean()
                                    viento_prom = df_power['viento_2m'].mean()
                                    
                                    # Asignar los mismos valores promedio a todas las zonas
                                    gdf_analizado['radiacion_solar'] = rad_prom
                                    gdf_analizado['viento_2m'] = viento_prom

                                    # --- Paso 2: Normalizar cada variable a [0, 1] ---
                                    def normalizar_solar(valor):
                                        return np.clip((valor - 3.0) / (7.0 - 3.0), 0, 1)

                                    def normalizar_viento(valor):
                                        return np.clip(1 - (valor - 1.0) / (5.0 - 1.0), 0, 1)

                                    def normalizar_humedad(ndwi):
                                        return np.clip((ndwi - 0.1) / (0.4 - 0.1), 0, 1)

                                    gdf_analizado['solar_norm'] = gdf_analizado['radiacion_solar'].apply(normalizar_solar)
                                    gdf_analizado['viento_norm'] = gdf_analizado['viento_2m'].apply(normalizar_viento)
                                    gdf_analizado['humedad_norm'] = gdf_analizado['ndwi'].apply(normalizar_humedad)

                                    # --- Paso 3: Calcular √≠ndice integrado ---
                                    w_fertilidad = 0.40
                                    w_solar = 0.25
                                    w_humedad = 0.20
                                    w_viento = 0.15

                                    gdf_analizado['potencial_cosecha'] = (
                                        w_fertilidad * gdf_analizado['npk_actual'] +
                                        w_solar * gdf_analizado['solar_norm'] +
                                        w_humedad * gdf_analizado['humedad_norm'] +
                                        w_viento * gdf_analizado['viento_norm']
                                    ).clip(0, 1)

                                    # Escalar a toneladas/ha seg√∫n cultivo base
                                    produccion_base = {
                                        'PALMA ACEITERA': 20,
                                        'CACAO': 1.2,
                                        'BANANO': 35,
                                        'CAF√â': 2.5
                                    }
                                    base = produccion_base.get(cultivo, 10)
                                    gdf_analizado['produccion_estimada'] = gdf_analizado['potencial_cosecha'] * base

                                    # --- Paso 4: Mostrar m√©tricas resumen ---
                                    col_c1, col_c2, col_c3, col_c4 = st.columns(4)
                                    with col_c1:
                                        st.metric("Potencial Promedio", f"{gdf_analizado['potencial_cosecha'].mean():.2f}")
                                    with col_c2:
                                        st.metric("M√°ximo", f"{gdf_analizado['potencial_cosecha'].max():.2f}")
                                    with col_c3:
                                        st.metric("Producci√≥n Estimada", f"{gdf_analizado['produccion_estimada'].mean():.1f} t/ha")
                                    with col_c4:
                                        total_est = (gdf_analizado['produccion_estimada'] * gdf_analizado['area_ha']).sum()
                                        st.metric("Total Parcela", f"{total_est:.1f} t")

                                    # --- Paso 5: Crear mapa de calor con ESRI ---
                                    mapa_calor = crear_mapa_potencial_cosecha_calor(gdf_analizado, cultivo)
                                    if mapa_calor:
                                        st.image(mapa_calor, use_container_width=True)
                                        st.download_button(
                                            "üì• Descargar Mapa de Calor",
                                            mapa_calor,
                                            f"potencial_cosecha_calor_{cultivo}_{datetime.now().strftime('%Y%m%d')}.png",
                                            "image/png"
                                        )

                                    # --- Paso 6: Identificar puntos calientes ---
                                    zonas_calientes = gdf_analizado[gdf_analizado['potencial_cosecha'] > 0.7]
                                    if not zonas_calientes.empty:
                                        st.subheader("üìç ZONAS CALIENTES (Potencial > 0.7)")
                                        st.dataframe(zonas_calientes[['id_zona', 'area_ha', 'potencial_cosecha', 'produccion_estimada']].sort_values('potencial_cosecha', ascending=False))
                                        
                                        total_area_caliente = zonas_calientes['area_ha'].sum()
                                        st.metric(f"√Årea total de zonas calientes", f"{total_area_caliente:.2f} ha")
                                        
                                        st.markdown("**Recomendaciones para zonas calientes:**")
                                        st.markdown("""
                                        - ‚úÖ **Maximizar inversi√≥n** en estas zonas (fertilizaci√≥n, riego)
                                        - ‚úÖ **Priorizar cosecha** temprana en estas √°reas
                                        - ‚úÖ **Monitoreo intensivo** para mantener altos rendimientos
                                        """)

                            # Crear mapa est√°tico con ESRI para an√°lisis GEE
                            if analisis_tipo in ["FERTILIDAD ACTUAL", "RECOMENDACIONES NPK"]:
                                columna_valor = 'valor_recomendado' if analisis_tipo == "RECOMENDACIONES NPK" else 'npk_actual'
                                mapa_buffer = crear_mapa_estatico_con_esri(gdf_analizado, 
                                                                          f"AN√ÅLISIS {analisis_tipo}", 
                                                                          columna_valor, 
                                                                          analisis_tipo, 
                                                                          nutriente, 
                                                                          cultivo, 
                                                                          satelite_seleccionado)
                                if mapa_buffer:
                                    st.subheader(f"üó∫Ô∏è MAPA CON ESRI SATELLITE - {analisis_tipo}")
                                    st.image(mapa_buffer, use_container_width=True)
                                    st.session_state['resultados_guardados']['mapa_buffer'] = mapa_buffer
                                    st.download_button(
                                        "üì• Descargar Mapa GEE con ESRI",
                                        mapa_buffer,
                                        f"mapa_gee_esri_{cultivo}_{satelite_seleccionado}_{analisis_tipo.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                                        "image/png"
                                    )
                                
                                st.subheader("üî¨ √çNDICES SATELITALES GEE POR ZONA")
                                columnas_indices = ['id_zona', 'npk_actual', 'materia_organica', 'ndvi', 'ndre', 'humedad_suelo', 'ndwi']
                                if analisis_tipo == "RECOMENDACIONES NPK":
                                    columnas_indices = ['id_zona', 'valor_recomendado', 'npk_actual', 'materia_organica', 'ndvi', 'ndre', 'humedad_suelo', 'ndwi']
                                columnas_indices = [col for col in columnas_indices if col in gdf_analizado.columns]
                                tabla_indices = gdf_analizado[columnas_indices].copy()
                                rename_dict = {
                                    'id_zona': 'Zona',
                                    'npk_actual': 'NPK Actual',
                                    'valor_recomendado': 'Recomendaci√≥n',
                                    'materia_organica': 'Materia Org (%)',
                                    'ndvi': 'NDVI',
                                    'ndre': 'NDRE',
                                    'humedad_suelo': 'Humedad',
                                    'ndwi': 'NDWI'
                                }
                                tabla_indices = tabla_indices.rename(columns={k: v for k, v in rename_dict.items() if k in tabla_indices.columns})
                                st.dataframe(tabla_indices)
        except Exception as e:
            st.error(f"‚ùå Error procesando archivo: {str(e)}")
            import traceback
            st.error(f"Detalle: {traceback.format_exc()}")
else:
    st.info("üìÅ Sube un archivo de tu parcela para comenzar el an√°lisis")

# ===== EXPORTACI√ìN PERSISTENTE =====
if 'resultados_guardados' in st.session_state:
    res = st.session_state['resultados_guardados']
    st.markdown("---")
    st.subheader("üì§ EXPORTAR RESULTADOS")
    col_exp1, col_exp2, col_exp3, col_exp4 = st.columns(4)
    with col_exp1:
        if st.button("üó∫Ô∏è Exportar GeoJSON", key="export_geojson"):
            geojson_data, nombre_archivo = exportar_a_geojson(res['gdf_analizado'], f"parcela_{res['cultivo']}")
            if geojson_data:
                st.download_button(
                    label="üì• Descargar GeoJSON",
                    data=geojson_data,
                    file_name=nombre_archivo,
                    mime="application/json",
                    key="geojson_download"
                )
    with col_exp2:
        if st.button("üìÑ Generar Reporte PDF", key="export_pdf"):
            with st.spinner("Generando PDF..."):
                estadisticas = generar_resumen_estadisticas(
                    res['gdf_analizado'],
                    res['analisis_tipo'],
                    res['cultivo'],
                    res.get('df_power')
                )
                recomendaciones = generar_recomendaciones_generales(res['gdf_analizado'], res['analisis_tipo'], res['cultivo'])
                mapa_buffer = res.get('mapa_buffer')
                pdf_buffer = generar_reporte_pdf(
                    res['gdf_analizado'], res['cultivo'], res['analisis_tipo'], res['area_total'],
                    res.get('nutriente'), res.get('satelite_seleccionado'), res.get('indice_seleccionado'),
                    mapa_buffer, estadisticas, recomendaciones
                )
                if pdf_buffer:
                    st.download_button(
                        label="üì• Descargar PDF",
                        data=pdf_buffer,
                        file_name=f"reporte_{res['cultivo']}_{res['analisis_tipo'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                        mime="application/pdf",
                        key="pdf_download"
                    )
                else:
                    st.error("‚ùå No se pudo generar el reporte PDF")
    with col_exp3:
        if st.button("üìù Generar Reporte DOCX", key="export_docx"):
            with st.spinner("Generando DOCX..."):
                estadisticas = generar_resumen_estadisticas(
                    res['gdf_analizado'],
                    res['analisis_tipo'],
                    res['cultivo'],
                    res.get('df_power')
                )
                recomendaciones = generar_recomendaciones_generales(res['gdf_analizado'], res['analisis_tipo'], res['cultivo'])
                mapa_buffer = res.get('mapa_buffer')
                docx_buffer = generar_reporte_docx(
                    res['gdf_analizado'], res['cultivo'], res['analisis_tipo'], res['area_total'],
                    res.get('nutriente'), res.get('satelite_seleccionado'), res.get('indice_seleccionado'),
                    mapa_buffer, estadisticas, recomendaciones
                )
                if docx_buffer:
                    st.download_button(
                        label="üì• Descargar DOCX",
                        data=docx_buffer,
                        file_name=f"reporte_{res['cultivo']}_{res['analisis_tipo'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="docx_download"
                    )
                else:
                    st.error("‚ùå No se pudo generar el reporte DOCX")
    with col_exp4:
        if st.button("üìä Exportar CSV", key="export_csv"):
            if res['gdf_analizado'] is not None:
                if 'geometry' in res['gdf_analizado'].columns:
                    df_export = res['gdf_analizado'].drop(columns=['geometry']).copy()
                else:
                    df_export = res['gdf_analizado'].copy()
                csv = df_export.to_csv(index=False)
                st.download_button(
                    label="üì• Descargar CSV",
                    data=csv,
                    file_name=f"datos_{res['cultivo']}_{res['analisis_tipo'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv",
                    key="csv_download"
                )

# FORMATOS ACEPTADOS Y METODOLOG√çA
with st.expander("üìã FORMATOS DE ARCHIVO ACEPTADOS"):
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**üó∫Ô∏è Shapefile (.zip)**")
        st.markdown("""
        - Archivo ZIP que contiene:
        - .shp (geometr√≠as)
        - .shx (√≠ndice)
        - .dbf (atributos)
        - .prj (proyecci√≥n, opcional)
        - Se recomienda usar EPSG:4326 (WGS84)
        """)
    with col2:
        st.markdown("**üåê KML (.kml)**")
        st.markdown("""
        - Formato Keyhole Markup Language
        - Usado por Google Earth
        - Contiene geometr√≠as y atributos
        - Puede incluir estilos y colores
        - Siempre en EPSG:4326
        """)
    with col3:
        st.markdown("**üì¶ KMZ (.kmz)**")
        st.markdown("""
        - Versi√≥n comprimida de KML
        - Archivo ZIP con extensi√≥n .kmz
        - Puede incluir recursos (im√°genes, etc.)
        - Compatible con Google Earth
        - Siempre en EPSG:4326
        """)

with st.expander("‚ÑπÔ∏è INFORMACI√ìN SOBRE LA METODOLOG√çA"):
    st.markdown("""
    **üå± SISTEMA DE AN√ÅLISIS MULTI-CULTIVO TROPICAL**
    **üõ∞Ô∏è SAT√âLITES SOPORTADOS:**
    - **Sentinel-2:** Alta resoluci√≥n (10m), revisita 5 d√≠as
    - **Landsat-8:** Resoluci√≥n media (30m), datos hist√≥ricos
    - **Datos Simulados:** Para pruebas y demostraciones
    **üìä CULTIVOS SOPORTADOS:**
    - **üå¥ PALMA ACEITERA:** Cultivo perenne con alta demanda de potasio
    - **üç´ CACAO:** Cultivo de sombra, requiere alta materia org√°nica
    - **üçå BANANO:** Cultivo exigente en nitr√≥geno y potasio, sensible a encharcamientos
    - **‚òï CAF√â:** Cultivo de monta√±a, sensible a pendientes y pH √°cido
    **üöÄ FUNCIONALIDADES MEJORADAS:**
    - **üå± Fertilidad Actual:** Estado NPK del suelo usando √≠ndices satelitales
    - **üó∫Ô∏è Mapas ESRI Satellite:** Mapas base de alta resoluci√≥n
    - **üî• Mapas de Calor:** Identificaci√≥n de puntos calientes para cosecha
    - **üíß NDWI (Humedad):** √çndice de Agua en Vegetaci√≥n/Suelo
    - **‚òÄÔ∏è Radiaci√≥n Solar:** Datos de NASA POWER (kWh/m¬≤/d√≠a)
    - **üí® Velocidad del Viento:** Datos de NASA POWER (m/s)
    - **üíß Precipitaci√≥n:** Datos de NASA POWER (mm/d√≠a)
    - **üíä Recomendaciones NPK:** Dosis espec√≠ficas por cultivo tropical
    - **üèóÔ∏è An√°lisis de Textura:** Composici√≥n del suelo (nomenclatura Venezuela/Colombia)
    - **üèîÔ∏è Curvas de Nivel:** An√°lisis topogr√°fico con mapa de calor de pendientes
    **üî¨ METODOLOG√çA CIENT√çFICA:**
    - An√°lisis basado en im√°genes satelitales
    - Integraci√≥n con datos meteorol√≥gicos de NASA POWER
    - Par√°metros espec√≠ficos para cultivos tropicales
    - C√°lculo de √≠ndices de vegetaci√≥n y suelo
    - Modelos digitales de elevaci√≥n (DEM) sint√©ticos
    - Recomendaciones validadas cient√≠ficamente
    **üí° CONSEJOS:**
    - Para mejores resultados, usa archivos en coordenadas EPSG:4326 (WGS84)
    - Los archivos KML deben contener pol√≠gonos (no puntos o l√≠neas)
    - El √°rea recomendada es entre 1 y 1000 hect√°reas
    - Todos los c√°lculos se realizan en EPSG:4326
    """)
